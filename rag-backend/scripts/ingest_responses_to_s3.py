"""
Ingest all .docx/.doc files from RAG_INFORMATION_DATABASE into FAISS and upload to S3.

Covers: Responses (875 docx), and any other non-PDF docs missed by the PDF ingestion run.

Run AFTER ingest_all_to_s3.py finishes. From rag-backend/:
  .\.venv_win\Scripts\python.exe scripts\ingest_responses_to_s3.py

Supports resume: checkpoint saved every 20 batches.
"""

import json
import logging
import os
import sys
import time
from pathlib import Path
from collections import defaultdict

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger(__name__)

SCRIPT_DIR = Path(__file__).resolve().parent
BASE_DIR   = SCRIPT_DIR.parent
sys.path.insert(0, str(BASE_DIR))

DB_ROOT     = BASE_DIR / "RAG_INFORMATION_DATABASE"
CHUNKS_FILE = BASE_DIR / "data" / "chunks" / "chunks.jsonl"
INDEX_FILE  = BASE_DIR / "vectordb" / "index.faiss"
META_FILE   = INDEX_FILE.with_suffix(".meta.json")

PENDING_FILE    = BASE_DIR / "vectordb" / "pending_docx_chunks.jsonl"
CHECKPOINT_FILE = BASE_DIR / "vectordb" / "ingest_docx_checkpoint.json"
INDEX_CKPT      = BASE_DIR / "vectordb" / "index.docx_checkpoint.faiss"
META_CKPT       = BASE_DIR / "vectordb" / "index.docx_checkpoint.meta.json"

S3_BUCKET     = "gst-rag-data-721082558531"
S3_REGION     = "ap-south-1"
S3_CHUNKS_KEY = "data/chunks/chunks.jsonl"
S3_INDEX_KEY  = "vectordb/index.faiss"
S3_META_KEY   = "vectordb/index.meta.json"

# All folders to scan for .docx/.doc files
CATEGORIES = [
    ("Responses",               "Response",    "responses"),
    ("Notification",            "Notification","notifications"),
    ("High Court Case Laws",    "Case Law",    "highcourt"),
    ("Supreme Court Case Laws", "Case Law",    "supremecourt"),
    ("Other APP Result",        "AAR",         "aar"),
    ("AAR",                     "AAR",         "aar"),
    ("Circulars",               "Circular",    "circulars"),
    ("Act",                     "Act",         "acts"),
    ("CGST",                    "CGST Act",    "cgst"),
    ("IGST",                    "IGST Act",    "igst"),
    ("Rules",                   "Rules",       "rules"),
    ("Forms",                   "Form",        "forms"),
    ("FAQs",                    "FAQ",         "faqs"),
    ("Brochures",               "Brochure",    "brochures"),
    ("ICAI",                    "ICAI",        "icai"),
]

# __MACOSX = Mac ZIP ghost folders — contain duplicate .docx stubs, never ingest
SKIP_FOLDERS   = {"generated_reports", "__MACOSX"}
CHECKPOINT_EVERY = 10


# ── S3 ─────────────────────────────────────────────────────────────────────────

def download_from_s3():
    import boto3
    s3 = boto3.client("s3", region_name=S3_REGION)
    for local, key in [(CHUNKS_FILE, S3_CHUNKS_KEY), (INDEX_FILE, S3_INDEX_KEY), (META_FILE, S3_META_KEY)]:
        local.parent.mkdir(parents=True, exist_ok=True)
        log.info(f"  {key} → {local.name}")
        s3.download_file(S3_BUCKET, key, str(local))
        log.info(f"    {local.stat().st_size / 1e6:.1f} MB")


def upload_to_s3():
    import boto3
    s3 = boto3.client("s3", region_name=S3_REGION)
    for local, key in [(CHUNKS_FILE, S3_CHUNKS_KEY), (INDEX_FILE, S3_INDEX_KEY), (META_FILE, S3_META_KEY)]:
        log.info(f"  {local.name} ({local.stat().st_size / 1e6:.1f} MB) → s3://{S3_BUCKET}/{key}")
        s3.upload_file(str(local), S3_BUCKET, key)


def restart_ecs():
    import boto3
    ecs = boto3.client("ecs", region_name=S3_REGION)
    ecs.update_service(cluster="gst-rag-cluster", service="gst-rag-backend-service", forceNewDeployment=True)
    log.info("  ECS force-redeploy triggered.")


# ── Index helpers ──────────────────────────────────────────────────────────────

def get_indexed_sources() -> set:
    indexed = set()
    if not CHUNKS_FILE.exists():
        return indexed
    with CHUNKS_FILE.open(encoding="utf-8", errors="replace") as f:
        for line in f:
            try:
                c = json.loads(line)
                src = (c.get("source") or c.get("rel_path") or
                       c.get("metadata", {}).get("source") or
                       c.get("metadata", {}).get("rel_path") or "")
                norm = src.replace("\\", "/").lower()
                if norm:
                    indexed.add(norm)
            except Exception:
                pass
    log.info(f"Existing indexed sources: {len(indexed)}")
    return indexed


def load_meta(path: Path) -> list:
    if not path.exists():
        return []
    with path.open(encoding="utf-8", errors="replace") as f:
        raw = f.read().strip()
    last = raw.rfind("]")
    if last != -1:
        raw = raw[:last + 1]
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return []


# ── Checkpointing ──────────────────────────────────────────────────────────────

def save_checkpoint(batch_idx, total, index, meta_list):
    import faiss
    faiss.write_index(index, str(INDEX_CKPT))
    with META_CKPT.open("w", encoding="utf-8") as f:
        json.dump(meta_list, f, ensure_ascii=False)
    CHECKPOINT_FILE.write_text(json.dumps({
        "batch_completed": batch_idx,
        "total_chunks": total,
        "saved_at": time.time(),
    }))
    log.info(f"  [CHECKPOINT] Saved after batch {batch_idx}")


def load_checkpoint():
    if not CHECKPOINT_FILE.exists():
        return None
    try:
        return json.loads(CHECKPOINT_FILE.read_text())
    except Exception:
        return None


def cleanup_checkpoint():
    for f in [PENDING_FILE, CHECKPOINT_FILE, INDEX_CKPT, META_CKPT]:
        if f.exists():
            f.unlink()


# ── Text extraction ────────────────────────────────────────────────────────────

def extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also grab table cells — Responses often store data in tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += "\n" + row_text
        return text.strip()
    except Exception as e:
        log.warning(f"  docx extraction failed for {path.name}: {e}")
        return ""


def ingest_docx(path: Path, document_type: str, category_key: str) -> list:
    import hashlib

    rel_path = str(path.relative_to(DB_ROOT)).replace("\\", "/")
    year = None
    for part in path.parts:
        if part.isdigit() and 2010 <= int(part) <= 2030:
            year = part
            break

    full_text = extract_docx(path)
    if not full_text or len(full_text) < 80:
        log.warning(f"  No usable text: {path.name}")
        return []

    WINDOW, STEP = 1500, 1300
    chunks = []
    for i, start in enumerate(range(0, max(1, len(full_text) - 200), STEP)):
        text = full_text[start: start + WINDOW].strip()
        if len(text) < 80:
            continue
        chunk_id = hashlib.md5(f"{rel_path}_{i}".encode()).hexdigest()[:16]
        meta = {
            "source": str(path), "rel_path": rel_path,
            "document_type": document_type, "category": category_key,
            "filename": path.name, "chunk_id": chunk_id,
            "file_type": path.suffix.lower().lstrip("."),
        }
        if year:
            meta["year"] = year
        chunks.append({"chunk_id": chunk_id, "text": text,
                        "source": str(path), "rel_path": rel_path, "metadata": meta})
    return chunks


# ── Embedding ──────────────────────────────────────────────────────────────────

def embed_and_append_all(all_new_chunks: list, resume_from_batch: int = 0):
    import faiss
    from sentence_transformers import SentenceTransformer

    log.info("Loading BAAI/bge-large-en-v1.5 model...")
    model = SentenceTransformer("BAAI/bge-large-en-v1.5")
    log.info("Model loaded.")

    if resume_from_batch > 0 and INDEX_CKPT.exists():
        log.info(f"Resuming from checkpoint (batch {resume_from_batch})...")
        index     = faiss.read_index(str(INDEX_CKPT))
        meta_list = load_meta(META_CKPT)
    else:
        index     = faiss.read_index(str(INDEX_FILE))
        meta_list = load_meta(META_FILE)

    log.info(f"FAISS base: {index.ntotal} vectors")

    BATCH = 64
    total = len(all_new_chunks)
    batches = (total + BATCH - 1) // BATCH

    for b_idx, start in enumerate(range(0, total, BATCH)):
        if b_idx < resume_from_batch:
            continue
        batch = all_new_chunks[start: start + BATCH]
        log.info(f"  Embedding batch {b_idx + 1}/{batches} ({len(batch)} chunks)...")
        embs = model.encode([c["text"] for c in batch], batch_size=32,
                            normalize_embeddings=True, show_progress_bar=False)
        embs = embs.astype("float32")
        index.add(embs)
        meta_list.extend([c["metadata"] for c in batch])

        if (b_idx + 1) % CHECKPOINT_EVERY == 0:
            save_checkpoint(b_idx, total, index, meta_list)

    log.info(f"FAISS after: {index.ntotal} vectors")
    faiss.write_index(index, str(INDEX_FILE))
    with META_FILE.open("w", encoding="utf-8") as f:
        json.dump(meta_list, f, ensure_ascii=False)
    log.info("Saved locally.")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    log.info("=" * 70)
    log.info("LETA DOCX INGESTION PIPELINE (Responses + all non-PDF docs)")
    log.info("=" * 70)

    checkpoint = load_checkpoint()
    resuming   = checkpoint is not None and PENDING_FILE.exists() and INDEX_CKPT.exists()

    if resuming:
        resume_from = checkpoint["batch_completed"] + 1
        log.info(f"\nRESUMING from batch {resume_from}")
        all_new_chunks = []
        with PENDING_FILE.open(encoding="utf-8", errors="replace") as f:
            for line in f:
                try:
                    all_new_chunks.append(json.loads(line))
                except Exception:
                    pass
        log.info(f"Loaded {len(all_new_chunks)} pending chunks.")
    else:
        resume_from = 0

        log.info("\n[1/4] Downloading latest S3 base files...")
        download_from_s3()

        log.info("\n[2/4] Scanning existing index...")
        indexed = get_indexed_sources()

        log.info("\n[3/4] Finding all unindexed .docx/.doc files...")
        missing_by_cat = defaultdict(list)
        for folder_name, doc_type, cat_key in CATEGORIES:
            cat_root = DB_ROOT / folder_name
            if not cat_root.exists():
                continue
            all_docs = [
                p for p in cat_root.rglob("*")
                if p.is_file()
                and p.suffix.lower() in {".docx", ".doc"}
                and not any(sk in p.parts for sk in SKIP_FOLDERS)
                and not p.name.startswith("._")   # skip Mac metadata stubs
            ]
            missing = [p for p in sorted(all_docs)
                       if str(p).replace("\\", "/").lower() not in indexed]
            if missing:
                missing_by_cat[folder_name] = (missing, doc_type, cat_key)
            log.info(f"  {folder_name:<30} docx/doc={len(all_docs):>4}  MISSING={len(missing):>4}")

        total_missing = sum(len(v[0]) for v in missing_by_cat.values())
        log.info(f"\n  Total missing docx/doc: {total_missing}")

        if total_missing == 0:
            log.info("All docx/doc files already indexed. Nothing to do.")
            return

        log.info(f"\n[4/4] Extracting and chunking {total_missing} files...")
        all_new_chunks = []
        grand_idx = 0
        for folder_name, (docs, doc_type, cat_key) in missing_by_cat.items():
            log.info(f"\n  === {folder_name} ({len(docs)} files) ===")
            for doc_path in docs:
                grand_idx += 1
                log.info(f"  [{grand_idx}/{total_missing}] {doc_path.relative_to(DB_ROOT)}")
                all_new_chunks.extend(ingest_docx(doc_path, doc_type, cat_key))

        log.info(f"\nTotal chunks extracted: {len(all_new_chunks)}")
        if not all_new_chunks:
            log.error("No chunks produced.")
            return

        # Save pending for resume
        with PENDING_FILE.open("w", encoding="utf-8") as f:
            for c in all_new_chunks:
                f.write(json.dumps(c, ensure_ascii=False) + "\n")

        # Append to chunks.jsonl
        with CHUNKS_FILE.open("a", encoding="utf-8") as f:
            for c in all_new_chunks:
                f.write(json.dumps(c, ensure_ascii=False) + "\n")
        log.info(f"Appended {len(all_new_chunks)} chunks to chunks.jsonl")

    log.info(f"\nEmbedding {len(all_new_chunks)} chunks (resume from batch {resume_from})...")
    embed_and_append_all(all_new_chunks, resume_from_batch=resume_from)

    log.info("\nUploading to S3...")
    upload_to_s3()

    log.info("\nRestarting ECS...")
    restart_ecs()

    cleanup_checkpoint()

    log.info("\n" + "=" * 70)
    log.info("DONE — docx/doc files indexed, ECS restarting.")
    log.info("=" * 70)


if __name__ == "__main__":
    main()
