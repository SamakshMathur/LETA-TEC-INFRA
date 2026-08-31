import json
import logging
import time
import uuid

from fastapi import FastAPI, File, UploadFile, Form
from pydantic import BaseModel
from typing import List, Any, Optional
from pathlib import Path

from app.routing.router import route_query
from app.generation.context_builder import build_context

# ── Structured JSON logging setup ────────────────────────────────────────────
# Each log line is a JSON object — queryable in CloudWatch Logs Insights.
# Fields: timestamp, level, logger, message, plus any extra kwargs passed
# to logger.info(..., extra={...}).

class _JsonFormatter(logging.Formatter):
    def format(self, record: logging.LogRecord) -> str:
        payload = {
            "timestamp": self.formatTime(record, "%Y-%m-%dT%H:%M:%S"),
            "level": record.levelname,
            "logger": record.name,
            "message": record.getMessage(),
        }
        # Attach any extra fields (query_id, latency_ms, cache_hit, etc.)
        for key, val in record.__dict__.items():
            if key not in (
                "name", "msg", "args", "levelname", "levelno", "pathname",
                "filename", "module", "exc_info", "exc_text", "stack_info",
                "lineno", "funcName", "created", "msecs", "relativeCreated",
                "thread", "threadName", "processName", "process", "message",
                "taskName",
            ):
                payload[key] = val
        return json.dumps(payload, default=str)


def _configure_json_logging():
    formatter = _JsonFormatter()
    root = logging.getLogger()
    if not root.handlers:
        handler = logging.StreamHandler()
        handler.setFormatter(formatter)
        root.addHandler(handler)
    else:
        for h in root.handlers:
            h.setFormatter(formatter)
    root.setLevel(logging.INFO)


def compute_context_fingerprint(
    question: str,
    chunks: list,
    assembled_context: str,
    provider: str,
    model: str
) -> dict:
    import hashlib
    import json

    chunk_identifiers = []
    chunk_hashes = []
    for idx, c in enumerate(chunks or []):
        c_path = c.get("rel_path") or c.get("metadata", {}).get("rel_path") or c.get("source") or ""
        c_page = c.get("page") or c.get("metadata", {}).get("page") or 0
        chunk_identifiers.append(f"{c_path}:{c_page}")

        c_text = c.get("text") or ""
        txt_hash = hashlib.sha256(c_text.encode("utf-8")).hexdigest()
        chunk_hashes.append({
            "index": idx,
            "identifier": f"{c_path}:{c_page}",
            "text_hash": txt_hash
        })

    context_hash = hashlib.sha256(assembled_context.encode("utf-8")).hexdigest()

    context_fingerprint_payload = {
        "question": question,
        "chunk_identifiers": chunk_identifiers,
        "chunk_hashes": chunk_hashes,
        "context_hash": context_hash
    }

    fingerprint_str = json.dumps(context_fingerprint_payload, sort_keys=True)
    fingerprint_hash = hashlib.sha256(fingerprint_str.encode("utf-8")).hexdigest()

    return {
        "fingerprint_hash": fingerprint_hash,
        "context_hash": context_hash,
        "chunk_identifiers": chunk_identifiers,
        "chunk_hashes": chunk_hashes,
        "question": question,
        "answer_provider": provider,
        "answer_model": model
    }


_configure_json_logging()
logger = logging.getLogger(__name__)

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # ─── Startup Tasks ───
    logger.info("LETA/Sentinel.AI starting up via Lifespan...")

    # 1. Validate Config
    try:
        from app.config import validate_config
        config_ok = validate_config()
        if not config_ok:
            logger.error("Configuration validation failed — check warnings")
        else:
            logger.info("Configuration validated successfully")
    except Exception as e:
        logger.error(f"Config validation error: {e}", exc_info=True)

    # 2. Run Vector Store Validation Gate
    try:
        from app.retrieval.retriever import validate_vector_store
        validate_vector_store()
    except Exception as e:
        logger.error(f"Vector store validation failed: {e}", exc_info=True)

    # 3. Seed feed store
    try:
        from app.api.documents import get_activity_feed
        from app.feed_store import _event_log
        items = get_activity_feed()
        for item in reversed(items):
            _event_log.append(item)
        logger.info(f"Feed store seeded with {len(items)} events from filesystem")
    except Exception as e:
        logger.warning(f"Feed store seed failed (non-fatal): {e}")

    # 4. Pre-load embedding model + FAISS index in the background
    import asyncio as _asyncio
    async def _warmup():
        try:
            from app.dependencies import preload_all_models
            await _asyncio.to_thread(preload_all_models)
            logger.info("Startup model warmup complete")
        except Exception as e:
            logger.warning(f"Startup warmup failed (non-fatal): {e}")

    _asyncio.ensure_future(_warmup())

    yield

    # ─── Shutdown Tasks ───
    logger.info("LETA/Sentinel.AI shutting down via Lifespan...")

    # 1. Close MongoDB Connection (if active)
    from app.database import db
    if db.client:
        try:
            logger.info("Closing MongoDB connection pool...")
            db.client.close()
            db.client = None
        except Exception as e:
            logger.warning(f"Error closing MongoDB client: {e}")

    # 2. Close Redis Cache Connection (if active)
    try:
        from app.cache import close_redis
        close_redis()
    except Exception as e:
        logger.warning(f"Error closing Redis client: {e}")


# ---------- App ----------
app = FastAPI(
    title="GST Legal RAG API",
    version="1.0",
    description="In-house GST knowledge assistant",
    lifespan=lifespan,
)


@app.get("/")
async def root():
    return {
        "message": "GST Legal RAG API is running.",
        "docs": "/docs",
        "health": "/api/health"
    }

from fastapi.middleware.cors import CORSMiddleware
import os

# CORS: Merge ALLOWED_ORIGINS env var with local dev origins so Vite
# preflight requests keep working even when production origins are configured.
_default_origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:5174",
    "http://127.0.0.1:5174",
    "https://gst-rag-95li.vercel.app",
    "https://main.d1q7i80dk455hq.amplifyapp.com",
    "https://letatec.com",
    "https://www.letatec.com",
]
_env_origins = [
    origin.strip()
    for origin in os.getenv("ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]
ALLOWED_ORIGINS = list(dict.fromkeys([*_default_origins, *_env_origins]))

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_origin_regex=r"^https://([a-z0-9-]+\.)?(vercel\.app|amplifyapp\.com)$",
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-Request-ID"],
    expose_headers=["X-Request-ID"],
    max_age=600,
)

# ── Security headers middleware ───────────────────────────────────────────────
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        response = await call_next(request)
        # Prevent browsers from MIME-sniffing the content type
        response.headers["X-Content-Type-Options"] = "nosniff"
        # Block clickjacking — skip for document view so the frontend iframe can embed PDFs
        if not request.url.path.startswith("/api/documents/view"):
            response.headers["X-Frame-Options"] = "DENY"
        # Force HTTPS for 1 year, include subdomains
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains; preload"
        # Only send origin in Referer header, no full URL
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        # Disable browser features not needed by the API
        response.headers["Permissions-Policy"] = "geolocation=(), camera=(), microphone=(), payment=()"
        # Prevent caching of API responses (contains legal/confidential data)
        if not request.url.path.startswith("/api/documents/view"):
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, private"
        for h in ("server", "x-powered-by"):
            if h in response.headers:
                del response.headers[h]
        return response

app.add_middleware(SecurityHeadersMiddleware)

# ── Request size limit middleware ─────────────────────────────────────────────
class RequestSizeLimitMiddleware(BaseHTTPMiddleware):
    MAX_BODY_BYTES = 20 * 1024 * 1024  # 20 MB

    async def dispatch(self, request: StarletteRequest, call_next):
        if request.method in ("POST", "PUT", "PATCH"):
            content_length = request.headers.get("content-length")
            if content_length and int(content_length) > self.MAX_BODY_BYTES:
                from starlette.responses import JSONResponse
                return JSONResponse({"detail": "Request body too large"}, status_code=413)
        return await call_next(request)

app.add_middleware(RequestSizeLimitMiddleware)

from fastapi import Request

# ── Request logging middleware ────────────────────────────────────────────────
# Emits one JSON log line per request with fields queryable in CloudWatch:
#   query_id, method, path, status_code, latency_ms
# The /ask endpoint adds extra fields (cache_hit, domain, etc.) via its own logs.

_request_logger = logging.getLogger("leta.request")

@app.middleware("http")
async def log_requests(request: Request, call_next):
    query_id = str(uuid.uuid4())[:8]
    request.state.query_id = query_id
    t0 = time.monotonic()

    response = await call_next(request)

    latency_ms = round((time.monotonic() - t0) * 1000, 1)
    _request_logger.info(
        f"{request.method} {request.url.path} {response.status_code}",
        extra={
            "query_id": query_id,
            "method": request.method,
            "path": request.url.path,
            "status_code": response.status_code,
            "latency_ms": latency_ms,
            "client_ip": request.client.host if request.client else "unknown",
        },
    )
    return response


# ── Security headers middleware ───────────────────────────────────────────────
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains; preload"
    # Skip X-Frame-Options for document view so the frontend iframe can embed PDFs
    if not request.url.path.startswith("/api/documents/view"):
        response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response

# Rate limiting — user-ID keyed, Redis-backed when available
import base64 as _b64
import json as _json_ratelimit
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded


def _rate_limit_key(request: Request) -> str:
    """Prefer JWT user-ID over IP so limits survive proxy/CDN hops."""
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        try:
            parts = auth.split(".")
            if len(parts) == 3:
                padded = parts[1] + "=" * (-len(parts[1]) % 4)
                payload = _json_ratelimit.loads(_b64.b64decode(padded))
                if "sub" in payload:
                    return f"user:{payload['sub']}"
        except Exception:
            pass
    forwarded = request.headers.get("X-Forwarded-For", "")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return get_remote_address(request)


def _build_limiter() -> Limiter:
    _redis_url = os.getenv("REDIS_URL", "redis://localhost:6379/0")
    try:
        import redis as _r
        _r.from_url(_redis_url, socket_connect_timeout=1).ping()
        lim = Limiter(key_func=_rate_limit_key, storage_uri=_redis_url)
        logger.info("Rate limiter: Redis-backed (distributed)")
        return lim
    except Exception as _e:
        logger.warning(f"Rate limiter: in-memory fallback (Redis unavailable: {_e})")
        return Limiter(key_func=_rate_limit_key)


limiter = _build_limiter()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

def _get_user_info_from_req(request: Request) -> tuple:
    user_id = "anonymous"
    username = None
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        try:
            token = auth_header.split(" ")[1]
            from app.security import verify_token
            payload = verify_token(token, "access")
            if payload:
                user_id = payload.get("sub", "anonymous")
                username = payload.get("sub")
        except Exception:
            pass
    return user_id, username

from app.api import documents
app.include_router(documents.router, prefix="/api/documents", tags=["Documents"])

from app.api import advisory
app.include_router(advisory.router, prefix="/api/advisory", tags=["Advisory"])

from app.api import sessions
app.include_router(sessions.router, prefix="/api/sessions", tags=["Sessions"])

from app.api import auth
app.include_router(auth.router, prefix="/api/auth", tags=["Auth"])

from app.api import templates
app.include_router(templates.router, prefix="/api/templates", tags=["Templates"])

from app.api import admin
app.include_router(admin.router, prefix="/api/admin", tags=["Admin"])

from app.api import knowledge
app.include_router(knowledge.router, prefix="/api/admin/knowledge", tags=["Admin Knowledge"])

from app.api import control_center
app.include_router(control_center.router, prefix="/api/admin/control-center", tags=["Admin Control Center"])

from app.api import feed
app.include_router(feed.router, prefix="/api/feed", tags=["Feed"])

from app.api import transcribe
app.include_router(transcribe.router, prefix="/api", tags=["Transcribe"])

from app.api import payments
app.include_router(payments.router, tags=["Payments"])

# ---------- Request / Response ----------
class QuestionRequest(BaseModel):
    question: str
    session_id: Optional[str] = None
    intent: Optional[str] = "general"
    provider: Optional[str] = None
    model: Optional[str] = None

class Source(BaseModel):
    source: str
    page: int

class AnswerResponse(BaseModel):
    answer: str
    confidence: float
    intent: str
    sources: List[Source]
    reasoning: Optional[Any] = None

# ---------- Validation Registry Helper (Fix 5) ----------
def validate_provider_and_model(provider: Optional[str] = None, model: Optional[str] = None) -> tuple[str, str]:
    """
    Validates requested provider and model against allowlist.
    Returns (actual_provider, actual_model).
    Raises HTTPException 400 if invalid (no silent fallback).
    """
    from app.config import LLM_PROVIDER_REGISTRY, ANSWER_LLM_PROVIDER, ANSWER_LLM_MODEL

    # Resolve default provider
    if not provider:
        provider = ANSWER_LLM_PROVIDER

    # Normalise aliases
    provider = provider.lower()
    if provider == "claude":
        provider = "anthropic"

    if provider not in LLM_PROVIDER_REGISTRY:
        _logger.error(f"[VALIDATION] Rejected unsupported provider: {provider}")
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported provider: {provider}. Allowed: {list(LLM_PROVIDER_REGISTRY.keys())}"
        )

    prov_info = LLM_PROVIDER_REGISTRY[provider]
    if not prov_info["enabled"]:
        _logger.error(f"[VALIDATION] Rejected disabled provider (missing API Key): {provider}")
        raise HTTPException(
            status_code=400,
            detail=f"Provider {provider} is currently disabled (missing credentials)."
        )

    # Resolve default model
    if not model:
        model = prov_info["default_model"]

    if model not in prov_info["models"]:
        _logger.error(f"[VALIDATION] Rejected unsupported model {model} for provider {provider}")
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported model: {model} for provider: {provider}. Allowed: {list(prov_info['models'].keys())}"
        )

    return provider, model


# ---------- Lazy Load Retriever ----------
from app.dependencies import get_retriever
from app.database import get_session_collection
from datetime import datetime

# ---------- Model Config Endpoint ----------
@app.get("/api/config/models")
async def get_available_models():
    """Returns safe metadata about allowed providers and models (Fix 7)."""
    from app.config import LLM_PROVIDER_REGISTRY
    safe_registry = {}
    for prov_id, prov_info in LLM_PROVIDER_REGISTRY.items():
        if prov_info["enabled"]:
            safe_registry[prov_id] = {
                "display_name": prov_info["display_name"],
                "default_model": prov_info["default_model"],
                "models": {
                    model_id: {"display_name": model_info["display_name"]}
                    for model_id, model_info in prov_info["models"].items()
                    if model_info["enabled"]
                }
            }
    return safe_registry


# ---------- Endpoint ----------
@app.get("/api/health")
async def health_check():
    health_status = {
        "status": "active",
        "timestamp": datetime.now().isoformat(),
        "systems": {
            "api": "ok",
            "database": "unknown",
            "retriever": "unknown"
        }
    }

    # Check Database
    try:
        import asyncio as _asyncio_health
        from app.database import get_db
        _db = get_db()
        if _db is not None:
            await _asyncio_health.to_thread(_db.command, "ping")
            health_status["systems"]["database"] = "connected"
    except Exception as e:
        health_status["systems"]["database"] = f"error: {str(e)}"
        health_status["status"] = "degraded"

    # Check Retriever
    try:
        retriever = get_retriever()
        if retriever is not None:
            health_status["systems"]["retriever"] = "initialized"
    except Exception as e:
        health_status["systems"]["retriever"] = f"error: {str(e)}"
        health_status["status"] = "degraded"

    return health_status

async def stream_and_save(
    generator,
    session_id,
    user_query,
    chunks=None,
    context="",
    truth_rules_text="",
    query_vec=None,
    is_draft_session: bool = False,
    answer_provider: str = None,
    answer_model: str = None,
    context_fingerprint: dict = None
):
    """
    Wrapper that streams the LLM response AND runs post-generation
    accuracy layers before saving to DB.
    Supports Strict Mode (buffered streaming + repair gate) and Standard Mode.
    """
    import logging
    _logger = logging.getLogger("stream_and_save")
    import asyncio
    import json
    full_answer = ""
    unique_sources = []

    # 1. Immediately emit metadata about retrieved sources (Perplexity-style transparency)
    if chunks:
        try:
            unique_sources = []
            seen_src = set()
            import urllib.parse as _urlparse
            for c in chunks:
                _dedup_rel = c.get("rel_path", "") or c.get("metadata", {}).get("rel_path", "") or c.get("source", "")
                src_key = (_dedup_rel, c.get("page", 0))
                if src_key not in seen_src:
                    seen_src.add(src_key)
                    _raw_src = c.get("source", "") or c.get("metadata", {}).get("source", "")
                    _rel_path = c.get("rel_path", "") or c.get("metadata", {}).get("rel_path", "")
                    # Use rel_path for display name — it holds the real filename.
                    _basename = os.path.basename(_rel_path) if _rel_path else os.path.basename(_raw_src)
                    _enc_path = _urlparse.quote(_rel_path.replace("\\", "/"), safe="") if _rel_path else ""
                    _enc_name = _urlparse.quote(_basename, safe="")
                    _url = (
                        f"/api/documents/view_by_path?path={_enc_path}"
                        if _enc_path else
                        f"/api/documents/view?category=all&filename={_enc_name}"
                    )
                    _snippet = (c.get("text") or "").strip()
                    unique_sources.append({
                        "title": _basename or "Document",
                        "page": c.get("page", 1),
                        "url": _url,
                        "rel_path": _rel_path,
                        "score": float(c.get("_rerank_score", 0)),
                        "snippet": _snippet[:800] if _snippet else "",
                    })
                if len(unique_sources) >= 20:  # collect more, then sort + cap
                    break

            # Sort by relevance score — the reranker already encodes legal
            # authority (30% weight) + semantic similarity (50%) + topic (20%).
            # Most relevant doc wins regardless of type; linkifyLegalRefs no
            # longer uses sources[0] as a fallback so AAR-as-first-link is gone.
            unique_sources.sort(key=lambda s: s.get("score", 0), reverse=True)
            unique_sources = unique_sources[:8]

            metadata_payload = {
                "type": "metadata",
                "sources": unique_sources
            }
            if answer_provider:
                metadata_payload["answer_provider"] = answer_provider
            if answer_model:
                metadata_payload["answer_model"] = answer_model
            if context_fingerprint:
                metadata_payload["context_fingerprint"] = context_fingerprint

            yield f"__METADATA__:{json.dumps(metadata_payload)}__END_METADATA__"
        except Exception as me:
            _logger.warning(f"Metadata emission failed: {me}")

    import time
    t_gen_start = time.monotonic()

    # Determine strict mode
    is_strict = is_draft_session or any(kw in user_query.lower() for kw in [
        "draft", "notice", "reply", "appeal", "submission", "advisory", "scn", "show cause"
    ])

    repair_triggered = False
    final_status = "SUCCESS_VERIFIED"

    try:
        if is_strict:
            # Strict Mode: Buffer and validate before streaming
            _logger.info("Strict Mode active: Buffering response stream for integrity gate...")
            initial_answer = ""
            last_status_time = time.monotonic()
            for chunk in generator:
                initial_answer += chunk
                current_time = time.monotonic()
                if current_time - last_status_time >= 1.5:
                    word_count = len(initial_answer.split())
                    yield f"__STATUS__:{json.dumps({'msg': f'Drafting reply... ({word_count} words written)'})}__END_STATUS__"
                    last_status_time = current_time

            yield f"__STATUS__:{json.dumps({'msg': 'Evaluating compliance checks...', 'status': 'VERIFICATION_PENDING'})}__END_STATUS__"

            from app.generation.calculation_engine import process_internal_calculations
            cleaned_answer, calculated_claims = process_internal_calculations(initial_answer, chunks or [])
            if calculated_claims:
                final_status = "SUCCESS_DERIVED"

            from app.generation.validator import validate_answer_integrity
            try:
                val_res = await asyncio.wait_for(
                    asyncio.to_thread(
                        validate_answer_integrity,
                        cleaned_answer,
                        chunks or [],
                        True,
                        user_query,
                        calculated_claims
                    ),
                    timeout=5.0
                )
            except asyncio.TimeoutError:
                _logger.warning("Integrity Gate validation timed out — serving as verification_pending")
                val_res = {
                    "is_valid": True,
                    "warnings": ["Integrity check timed out. Verification is pending."],
                    "citations_status": {},
                    "ungrounded_numbers": [],
                    "severity": "NONE",
                    "degraded_fallback": True,
                    "timed_out": True
                }
                final_status = "VERIFICATION_PENDING"

            if val_res["is_valid"]:
                full_answer = cleaned_answer
                yield f"__STATUS__:{json.dumps({'msg': 'Compliance check complete.', 'status': final_status})}__END_STATUS__"
            else:
                # Execution of ONE automatic repair attempt
                _logger.warning(f"Integrity Gate failed ({val_res['severity']}). Executing repair loop...")
                yield f"__STATUS__:{json.dumps({'msg': 'Repairing compliance errors...', 'status': 'VERIFICATION_PENDING'})}__END_STATUS__"
                repair_triggered = True
                warnings_str = "\n".join(f"- {w}" for w in val_res["warnings"])

                correction_prompt = f"""
                [COMPLIANCE CORRECTION REQUIRED]
                Your previous answer failed legal compliance checks with these errors:
                {warnings_str}

                Please regenerate the response correcting these errors. Only cite valid facts from sources:
                Sources:
                {context[:8000]}

                Previous Answer:
                {initial_answer[:4000]}
                """

                from app.generation.synthesizer import synthesize_answer_stream as _synth_stream
                repaired_answer = ""
                repair_stream = _synth_stream(correction_prompt, context, session_is_draft=is_draft_session, provider=answer_provider, model=answer_model)
                last_status_time = time.monotonic()
                for chunk in repair_stream:
                    repaired_answer += chunk
                    current_time = time.monotonic()
                    if current_time - last_status_time >= 1.5:
                        word_count = len(repaired_answer.split())
                        yield f"__STATUS__:{json.dumps({'msg': f'Drafting compliance repair... ({word_count} words written)'})}__END_STATUS__"
                        last_status_time = current_time

                cleaned_repaired, calculated_claims_repaired = process_internal_calculations(repaired_answer, chunks or [])
                if calculated_claims_repaired:
                    final_status = "SUCCESS_DERIVED"

                # Re-validate the repaired answer
                try:
                    val_res_repaired = await asyncio.wait_for(
                        asyncio.to_thread(
                            validate_answer_integrity,
                            cleaned_repaired,
                            chunks or [],
                            True,
                            user_query,
                            calculated_claims_repaired
                        ),
                        timeout=5.0
                    )
                except asyncio.TimeoutError:
                    _logger.warning("Integrity Gate repair validation timed out — serving as verification_pending")
                    val_res_repaired = {
                        "is_valid": True,
                        "warnings": ["Repair check timed out. Verification is pending."],
                        "citations_status": {},
                        "ungrounded_numbers": [],
                        "severity": "NONE",
                        "degraded_fallback": True,
                        "timed_out": True
                    }
                    final_status = "VERIFICATION_PENDING"

                if val_res_repaired["is_valid"]:
                    _logger.info("Integrity Gate: Repair successful.")
                    full_answer = cleaned_repaired
                    yield f"__STATUS__:{json.dumps({'msg': 'Compliance check complete.', 'status': final_status})}__END_STATUS__"
                else:
                    # Second failure -> Return safe qualified fallback
                    _logger.error("Integrity Gate: Repair failed validation. Serving fallback disclaimer.")
                    final_status = "VERIFICATION_FAILED"
                    full_answer = (
                        "I cannot sufficiently verify this answer from the available legal sources.\n\n"
                        "**Unverified details:**\n"
                        f"{warnings_str}\n\n"
                        "Please review the source documents manually."
                    )
                    yield f"__STATUS__:{json.dumps({'msg': 'Compliance check failed.', 'status': final_status})}__END_STATUS__"

            # Stream the buffered/verified result to user with minor pacing
            for i in range(0, len(full_answer), 40):
                yield full_answer[i:i+40]
                await asyncio.sleep(0.005)

        else:
            # Standard Mode: Stream immediately to client, validate post-stream
            for chunk in generator:
                full_answer += chunk
                yield chunk

            # Post-stream compliance check
            try:
                from app.generation.calculation_engine import process_internal_calculations
                cleaned_answer, calculated_claims = process_internal_calculations(full_answer, chunks or [])
                full_answer = cleaned_answer

                from app.generation.validator import validate_answer_integrity, validate_logic
                val_res = await asyncio.to_thread(
                    validate_answer_integrity,
                    full_answer,
                    chunks or [],
                    False,
                    user_query,
                    calculated_claims
                )
                final_severity = val_res.get("severity", "NONE")
                all_warnings = val_res["warnings"] + validate_logic(full_answer)
                if all_warnings:
                    all_warnings = list(set(all_warnings))
                    warning_msg = "\n\n> [!WARNING] **AUTOMATED COMPLIANCE CHECK**\n"
                    warning_msg += "> The following potential issues were detected in this drafted opinion:\n"
                    for w in all_warnings:
                        warning_msg += f"> - {w}\n"
                    warning_msg += "> \n> *Please verify these points manually before professional use.*"
                    yield warning_msg
                    full_answer += warning_msg
                    final_status = "SUCCESS_PARTIALLY_VERIFIED"
                else:
                    final_status = "SUCCESS_DERIVED" if calculated_claims else "SUCCESS_VERIFIED"
            except Exception as val_err:
                _logger.warning(f"Post-stream validation failed: {val_err}")
                final_status = "VERIFICATION_PENDING"
                final_severity = "ERROR"

        if chunks and full_answer.strip():
            # Use the session-level draft flag passed in — do not re-detect from
            # user_query alone, which breaks on follow-up/correction messages.
            is_draft = is_draft_session or any(kw in user_query.lower() for kw in [
                "draft","notice","reply","appeal","submission","advisory","scn",
                "show cause","drc-01","drc 01","asmt-10","our understanding",
                "gst implications","provide opinion","our comments","tax position",
                "advise on","legal opinion","our client is","we are engaged in",
            ])

            if not is_draft:
                # --- Parallel Accuracy Pipeline ---
                async def run_citation_validator():
                    try:
                        from app.generation.citation_validator import CitationValidator
                        return CitationValidator.validate_citations(full_answer, chunks)
                    except Exception as e:
                        _logger.warning(f"Citation validator error: {e}")
                        return full_answer

                async def run_answer_verifier():
                    try:
                        from app.generation.answer_verifier import verify_answer
                        return await asyncio.to_thread(verify_answer, user_query, full_answer, chunks)
                    except Exception as e:
                        _logger.warning(f"Answer verifier error: {e}")
                        return None

                async def run_hallucination_guard():
                    try:
                        from app.generation.hallucination_guard import check_hallucinated_numbers
                        return check_hallucinated_numbers(full_answer, context, truth_rules_text, chunks)
                    except Exception as e:
                        _logger.warning(f"Hallucination guard error: {e}")
                        return ""

                async def run_template_matcher():
                    try:
                        from app.retrieval.template_matcher import search_templates, format_template_suggestions
                        matched = await asyncio.to_thread(search_templates, user_query, top_k=3)
                        return format_template_suggestions(matched)
                    except Exception as e:
                        _logger.warning(f"Template matcher error: {e}")
                        return ""

                try:
                    results = await asyncio.wait_for(
                        asyncio.gather(
                            run_citation_validator(),
                            run_hallucination_guard(),
                        ),
                        timeout=3.0,
                    )
                except asyncio.TimeoutError:
                    _logger.warning("Post-generation validators timed out — skipping")
                    results = (full_answer, "")

                _, hallu_warn = results

                # Log internally — validator output is NEVER streamed to the user.
                if hallu_warn:
                    _logger.warning(f"Hallucination guard: {hallu_warn[:300]}")

    except Exception as e:
        _logger.error(f"Error in stream_and_save: {e}", exc_info=True)
        yield "\n[System: An error occurred processing your request. Please try again.]"
    finally:
        # DB save — wrapped in to_thread so the sync pymongo write doesn't
        # block the event loop while the HTTP response is still open.
        if session_id and full_answer.strip():
            collection = get_session_collection()
            if collection is not None:
                try:
                    await asyncio.to_thread(
                        collection.update_one,
                        {"session_id": session_id},
                        {
                            "$push": {"messages": {
                                "message_id": str(uuid.uuid4()),
                                "role": "assistant",
                                "content": full_answer,
                                "metadata": {"status": final_status},
                                "citations": unique_sources,
                                "sources": unique_sources,
                                "timestamp": datetime.now(),
                            }},
                            "$set": {"updated_at": datetime.now()},
                        },
                    )
                except Exception as dbe:
                    _logger.warning(f"DB save error (non-fatal): {dbe}")

        # Confidence-gated cache store — also off the event loop
        if full_answer.strip() and query_vec is not None:
            try:
                from app.cache import cache_store
                from app.generation.confidence import estimate_confidence
                confidence = await asyncio.to_thread(estimate_confidence, full_answer, chunks or [])
                await asyncio.to_thread(cache_store, user_query, query_vec, full_answer, confidence, unique_sources if chunks else [])
                _logger.debug(f"Cache store attempt | confidence={confidence:.2f}")
            except Exception as ce:
                _logger.warning(f"Cache store error (non-fatal): {ce}")

        # Commit AI log analytics
        try:
            from app.ai_logger import update_ai_log, commit_ai_log
            if "val_res_repaired" in locals():
                gate_severity = val_res_repaired.get("severity", "NONE")
            elif "val_res" in locals():
                gate_severity = val_res.get("severity", "NONE")
            else:
                gate_severity = "NONE"

            update_ai_log(
                generation_time_ms=round((time.monotonic() - t_gen_start) * 1000, 2),
                estimated_completion_tokens=len(full_answer) // 4,
                response_length=len(full_answer),
                citations_count=len(unique_sources),
                integrity_gate_result=gate_severity,
                repair_triggered=repair_triggered
            )
            commit_ai_log(success=bool(full_answer.strip()))
        except Exception as cle:
            _logger.warning(f"AI logger commit error (non-fatal): {cle}")

def check_and_escalate_evidence(question: str, chunks: list, allowed_sources=None, domain_paths=None, is_draft: bool = False) -> tuple[list, bool, list]:
    """
    Verifies if all explicit citations mentioned in the query exist in the chunks.
    If any citation is missing, triggers Retrieval Escalation:
    1. Force-lookup the exact provision index for the missing reference key.
    2. Add matched canonical primary/secondary chunks to the result pool.
    Returns (chunks, coverage_satisfied, list_of_still_missing_provisions).
    """
    from app.retrieval.reference_resolver import ReferenceResolver
    from app.retrieval.retriever import _is_quarantined

    resolved = ReferenceResolver.resolve_references(question)
    if not resolved:
        return chunks, True, []  # no explicit references, trivially satisfied

    retriever = get_retriever()
    current_chunk_ids = {c.get("chunk_id") for c in chunks}

    missing_refs = []
    for ref_obj in resolved:
        key = ref_obj["canonical_key"]

        # Check if this key exists in the retrieved chunks provisions/citations
        found = False
        for c in chunks:
            c_refs = c.get("metadata", {}).get("provisions", []) + c.get("metadata", {}).get("citations", []) + c.get("provisions", [])
            if any(key in str(r) for r in c_refs):
                found = True
                break

        if not found:
            missing_refs.append(ref_obj)

    if not missing_refs:
        return chunks, True, []  # all references covered!

    # Trigger Retrieval Escalation
    _logger = logging.getLogger("leta.escalation")
    _logger.info(f"Retrieval Escalation triggered for missing citations: {[r['canonical_key'] for r in missing_refs]}")
    escalated_chunks = list(chunks)

    for ref_obj in missing_refs:
        key = ref_obj["canonical_key"]

        # Directly fetch all chunks matching the provision key from the provision index
        indices = retriever._provision_index.get(key, [])
        if key.startswith("CIRCULAR_") and hasattr(retriever, "_circular_index"):
            indices = list(set(indices + retriever._circular_index.get(key, [])))

        ref_chunks = []
        for idx in indices:
            if idx >= len(retriever.chunks):
                continue
            c = retriever.chunks[idx]
            if _is_quarantined(c):
                continue
            rel_path = c.get("rel_path") or c.get("metadata", {}).get("rel_path", "")
            if rel_path in retriever.inactive_paths:
                continue

            ref_chunks.append(c)

        if not ref_chunks:
            continue

        # Sort the direct matches: primary Act/Rule matches first
        primary_matches = []
        secondary_matches = []
        for c in ref_chunks:
            meta = c.get("metadata", {})
            doc_type = meta.get("canonical_document_type", meta.get("document_type", "REFERENCE")).upper()
            rel_path_lower = (c.get("rel_path") or meta.get("rel_path", "")).lower()

            is_primary = False
            filename = os.path.basename(rel_path_lower)
            in_acts_folder = any(k in rel_path_lower for k in ["acts/", "act/", "cgst acts", "igst acts"])
            in_rules_folder = any(k in rel_path_lower for k in ["rules/", "rule/", "cgst rules", "igst rules"])

            if "_SEC_" in key:
                sec_part = key.split("_SEC_")[-1]
                m_num = re.search(r'\d+', sec_part)
                if m_num:
                    target_sec = int(m_num.group(0))
                    range_match = re.search(r'\b(?:sections?|secs?\.?)\s*(\d+)(?:\s*[\–\-]\s*(\d+))?', filename, re.IGNORECASE)
                    if range_match:
                        start = int(range_match.group(1))
                        end = int(range_match.group(2)) if range_match.group(2) else start
                        is_primary = (start <= target_sec <= end)
                    else:
                        is_primary = in_acts_folder or (doc_type == "PRIMARY_LAW" and any(k in rel_path_lower for k in ["acts", "act"]))
                else:
                    is_primary = in_acts_folder or (doc_type == "PRIMARY_LAW" and any(k in rel_path_lower for k in ["acts", "act"]))
            elif "_RUL_" in key:
                rule_part = key.split("_RUL_")[-1]
                if in_rules_folder and rule_part.lower() in filename:
                    is_primary = True
                else:
                    is_primary = in_rules_folder or (doc_type == "RULES" and any(k in rel_path_lower for k in ["rules", "rule"]))
            elif "_NOT_" in key:
                is_primary = (doc_type == "NOTIFICATION" or "notification" in rel_path_lower)
            elif "CIRCULAR_" in key:
                is_primary = (doc_type == "CIRCULAR" or "circular" in rel_path_lower)

            if is_primary:
                primary_matches.append(c)
            else:
                secondary_matches.append(c)

        # Dynamic allocation for escalated items
        added = 0
        for c in primary_matches:
            cid = c.get("chunk_id")
            if cid and cid not in current_chunk_ids:
                c_copy = c.copy()
                c_copy["_pinned_by_ref"] = True
                c_copy["_pinned_tier"] = "PRIMARY"
                c_copy["_statute_priority"] = 1.0
                escalated_chunks.append(c_copy)
                current_chunk_ids.add(cid)
                added += 1
                if added >= 3:
                    break

        for c in secondary_matches:
            if added >= 5:
                break
            cid = c.get("chunk_id")
            if cid and cid not in current_chunk_ids:
                c_copy = c.copy()
                c_copy["_pinned_by_ref"] = True
                c_copy["_pinned_tier"] = "SECONDARY"
                c_copy["_statute_priority"] = 0.8
                escalated_chunks.append(c_copy)
                current_chunk_ids.add(cid)
                added += 1

    # Verify coverage again after escalation
    still_missing = []
    for ref_obj in resolved:
        key = ref_obj["canonical_key"]
        found = False
        for c in escalated_chunks:
            c_refs = c.get("metadata", {}).get("provisions", []) + c.get("metadata", {}).get("citations", []) + c.get("provisions", [])
            if any(key in str(r) for r in c_refs):
                found = True
                break
        if not found:
            still_missing.append(ref_obj["provision"])

    if still_missing:
        _logger.warning(f"Retrieval Escalation failed to recover evidence for: {still_missing}")
        return escalated_chunks, False, still_missing

    return escalated_chunks, True, []


_ask_logger = logging.getLogger("leta.ask")

@app.post("/ask")
@limiter.limit("30/minute")
async def ask_question(request: Request, req: QuestionRequest):
    question = req.question.strip()
    session_id = req.session_id
    query_id = getattr(request.state, "query_id", str(uuid.uuid4())[:8])
    t0 = time.monotonic()

    # 1. Validate requested provider & model (Fix 5)
    actual_provider, actual_model = validate_provider_and_model(req.provider, req.model)

    user_id, username = _get_user_info_from_req(request)
    from app.ai_logger import init_ai_log
    init_ai_log(
        user_id=user_id,
        username=username,
        query=question,
        endpoint="/api/ask",
        session_id=session_id,
        request_id=query_id,
        client_ip=request.client.host if request.client else None
    )
    from app.ai_logger import update_ai_log
    update_ai_log(
        selected_provider=req.provider or "default",
        selected_model=req.model or "default",
        provider=actual_provider,
        model_used=actual_model
    )

    # IMMEDIATE SAVE: Save User Question First (off event loop)
    if session_id:
        _msg_collection = get_session_collection()
        if _msg_collection is not None:
            import asyncio as _asyncio_save
            _user_msg = {
                "message_id": str(uuid.uuid4()),
                "role": "user",
                "content": question,
                "metadata": {},
                "citations": [],
                "sources": [],
                "timestamp": datetime.now(),
            }
            await _asyncio_save.to_thread(
                _msg_collection.update_one,
                {"session_id": session_id},
                {"$push": {"messages": _user_msg}},
            )

    async def rag_pipeline_orchestrator():
        import asyncio as _asyncio
        from app.routing.router import route_query
        from app.generation.synthesizer import _estimate_complexity
        from app.generation.calculation_engine import detect_and_calculate, format_for_context
        from app.generation.context_compressor import compress_context
        from app.cache import cache_lookup
        from app.retrieval.retriever import embed_query

        # ── Stage 1: Instant intent classification (pure keyword, ~1ms) ──────────
        route = route_query(question)
        domain_paths = route.get("domain_paths", [])
        _complexity = _estimate_complexity(question)
        _q = question.lower()

        _DRAFT_KW = [
            "draft","notice","reply","appeal","submission","advisory","scn","show cause",
            "drc-01","drc 01","asmt-10","asmt 10","drc-07","drc 07","drc-03","drc 03",
            "write a letter","write letter","prepare reply","representation",
            "response to notice","respond to","our understanding","gst implications",
            "gst implication","provide opinion","provide advisory","our comments",
            "tax position","gst treatment of","advise on","legal opinion",
            "our client is","we are engaged in","facts of the case",
        ]
        # Definition/explanation queries must NEVER route to DRAFTING_PROMPT regardless
        # of session history. A plain knowledge query ("define X", "what is X",
        # "provide definition of X") has no missing facts and should never trigger
        # the check-3 clarification flow.
        _NEVER_DRAFT_PATTERNS = [
            r'\b(what\s+is|what\s+are|define|definition\s+of|explain|meaning\s+of)\b',
            r'\b(provide|give|state|share)\s+(the\s+)?(definition|meaning|explanation|rate|provision)',
            r'\b(relevant\s+circular|applicable\s+circular|circular\s+on)\b',
            r'\b(full\s+form|abbreviation|what\s+does\s+.+stand\s+for)\b',
        ]
        import re as _re
        _is_never_draft = any(_re.search(p, _q) for p in _NEVER_DRAFT_PATTERNS)

        _is_draft_early = (not _is_never_draft) and any(k in _q for k in _DRAFT_KW)

        if "rule 42" in _q or "rule42" in _q:
            _init_msg = "Computing Rule 42 ITC Reversal..."
        elif "rule 43" in _q or "rule43" in _q:
            _init_msg = "Computing Rule 43 Capital Goods Reversal..."
        elif any(k in _q for k in ["itc", "input tax credit", "section 16", "section 17"]):
            _init_msg = "Analyzing ITC Eligibility Provisions..."
        elif _is_draft_early:
            _init_msg = "Initializing Advisory & Drafting Engine..."
        elif any(k in _q for k in ["refund", "export", "lut"]):
            _init_msg = "Checking Refund & Export Provisions..."
        elif any(k in _q for k in ["interest", "section 50", "penalty"]):
            _init_msg = "Computing Interest & Penalty Exposure..."
        else:
            _init_msg = "Initializing Statutory Analyzer..."

        yield f"__STATUS__:{json.dumps({'msg': _init_msg})}__END_STATUS__"

        # ── Stage 2: PARALLEL — session history + query embedding ─────────────────
        # Both are independent I/O-bound tasks; run them simultaneously.
        _HISTORY_DRAFT_KW = [
            "advisory","our understanding","gst implications","gst implication",
            "provide opinion","our comments","tax position","gst treatment",
            "advise on","legal opinion","our client","we are engaged",
            "facts of the case","draft","notice","reply","appeal","scn","show cause",
            "drc-01","drc-07","asmt-10","representation",
            "our comments from gst perspective","i've reviewed your",
            "issue raised","section invoked","to draft a strong reply",
        ]

        def _fetch_history_sync():
            if not session_id:
                return "", False
            _coll = get_session_collection()
            if _coll is None:
                return "", False
            _sess = _coll.find_one({"session_id": session_id})
            if not _sess or "messages" not in _sess:
                return "", False
            # IMPORTANT: include the current user message (already saved by the
            # immediate save above) in the history window. Previously [:-1] excluded
            # it, so CHECK 1 in the drafting prompt could not see the user's reply
            # to LETA's questions and would ask again. Using [-7:] shows the model
            # the full sequence: LETA-questions → USER-answers → produce draft.
            _recent = _sess["messages"][-7:]
            _hist = "".join(f"{m['role'].upper()}: {m['content']}\n" for m in _recent)
            _is_d = any(k in _hist.lower() for k in _HISTORY_DRAFT_KW)
            return _hist, _is_d

        yield f"__STATUS__:{json.dumps({'msg': 'Scanning Semantic Cache...'})}__END_STATUS__"

        (history_context, _session_is_draft), query_vec = await _asyncio.gather(
            _asyncio.to_thread(_fetch_history_sync),
            _asyncio.to_thread(embed_query, question),
        )
        _is_draft = _is_draft_early or _session_is_draft

        # ── Stage 3: Cache lookup ─────────────────────────────────────────────────
        cached_answer = await _asyncio.to_thread(cache_lookup, question, query_vec, provider=actual_provider, model=actual_model)
        if cached_answer:
            cached_text, cached_sources = cached_answer
            _ask_logger.info("Cache HIT", extra={"query_id": query_id, "cache_hit": True})
            yield f"__STATUS__:{json.dumps({'msg': 'Cache Hit — Instant Retrieval Complete.'})}__END_STATUS__"
            if cached_sources:
                yield f"__METADATA__:{json.dumps({'type': 'metadata', 'sources': cached_sources})}__END_METADATA__"
            yield cached_text
            try:
                from app.ai_logger import update_ai_log, commit_ai_log
                update_ai_log(
                    model_used="cache",
                    cache_hit=True,
                    response_length=len(cached_text),
                    citations_count=len(cached_sources or [])
                )
                commit_ai_log(success=True)
            except Exception:
                pass
            return

        # ── Stage 4: Pre-computed calculations (sync, instant) ────────────────────
        calc_result = detect_and_calculate(question)
        calc_block = format_for_context(calc_result) if calc_result else ""
        if calc_result:
            yield f"__STATUS__:{json.dumps({'msg': 'Pre-computing Statutory Formula...'})}__END_STATUS__"

        # ── Stage 5: Retrieval — strategy depends on query type ───────────────────
        # Follow-up query enrichment: if query is short (<8 words) and history
        # exists, enrich the retrieval query with context from the last exchange
        # so "what about penalties?" searches with full legal context.
        retriever = get_retriever()
        _retrieval_top_k = 15 if _is_draft else (12 if _complexity >= 0.60 else 8)
        domain_label = ", ".join(domain_paths[:2]) if domain_paths else "All Databases"

        def _enrich_for_retrieval(q: str, hist: str) -> str:
            if hist and len(q.split()) < 8:
                _last_lines = [l for l in hist.split('\n') if l.startswith('ASSISTANT:')]
                if _last_lines:
                    _ctx = _last_lines[-1].replace('ASSISTANT:', '').strip()[:100]
                    return f"{q} [context: {_ctx}]"
            return q

        _retrieval_q = _enrich_for_retrieval(question, history_context)

        yield f"__STATUS__:{json.dumps({'msg': f'Searching Statutory Database ({domain_label})...'})}__END_STATUS__"

        # Rule-based query topic routing
        from app.retrieval.query_refiner import classify_topic_rules, extract_query_topic
        topic_info = classify_topic_rules(question)
        if topic_info.get("topic") is None:
            topic_info = await _asyncio.to_thread(extract_query_topic, question, provider=actual_provider, model=actual_model)

        if _is_draft:
            # Draft: rule-based sub-queries, no LLM expansion call needed
            _statute_q  = _retrieval_q + " section rule act provisions conditions eligibility liability"
            _caselaw_q  = _retrieval_q + " high court supreme court judgment held ruling decision AAR"
            _circular_q = _retrieval_q + " CBIC circular notification clarification instruction"
            _adv = {
                "queries": [_retrieval_q, _statute_q, _caselaw_q, _circular_q],
                "hyde_document": "",
                "topic": topic_info.get("topic") or "General",
                "subtopic": topic_info.get("subtopic"),
            }
            chunks = await _asyncio.to_thread(
                retriever.search, _retrieval_q, _retrieval_top_k,
                route["use_sources"], _adv, domain_paths, True,
            )

        elif _complexity >= 0.35:
            # Complex non-draft: run query expansion + fast retrieval IN PARALLEL.
            # Fast retrieval uses original query without expansion (skip_rerank for speed).
            # When expansion arrives, we supplement the pool and do ONE final rerank.
            yield f"__STATUS__:{json.dumps({'msg': 'Expanding Query for Precision Retrieval...'})}__END_STATUS__"

            def _fast_retrieve():
                return retriever.search(
                    _retrieval_q, _retrieval_top_k, route["use_sources"],
                    None, domain_paths, False, skip_rerank=True,
                )

            def _expand():
                from app.retrieval.query_refiner import generate_advanced_queries
                return generate_advanced_queries(_retrieval_q, provider=actual_provider, model=actual_model)

            fast_chunks, advanced_queries = await _asyncio.gather(
                _asyncio.to_thread(_fast_retrieve),
                _asyncio.to_thread(_expand),
            )

            # Override with rule-based topic if matched
            if topic_info.get("topic") is not None:
                advanced_queries["topic"] = topic_info["topic"]
                advanced_queries["subtopic"] = topic_info["subtopic"]

            # Supplement the fast pool with expanded-query FAISS results,
            # then do ONE FlashRank + LegalReranker pass on the merged pool.
            chunks = await _asyncio.to_thread(
                retriever.supplement_and_rerank,
                fast_chunks, advanced_queries, _retrieval_q, _retrieval_top_k,
            )

        else:
            # Simple query: direct retrieval (already fast)
            _adv = {
                "queries": [_retrieval_q],
                "hyde_document": "",
                "topic": topic_info.get("topic") or "General",
                "subtopic": topic_info.get("subtopic")
            }
            chunks = await _asyncio.to_thread(
                retriever.search, _retrieval_q, _retrieval_top_k,
                route["use_sources"], _adv, domain_paths, False,
            )

        # Check and escalate retrieval coverage
        chunks, coverage_satisfied, still_missing = check_and_escalate_evidence(
            question, chunks, allowed_sources=route["use_sources"],
            domain_paths=domain_paths, is_draft=_is_draft
        )

        if not coverage_satisfied:
            yield f"__STATUS__:{json.dumps({'msg': 'Retrieval Escalation failed to recover required evidence.'})}__END_STATUS__"
            missing_str = ", ".join(still_missing)
            yield f"\nI cannot sufficiently answer your query because the required statutory evidence for the following provision(s) is missing from the database: **{missing_str}**. Please review the source documents manually."
            try:
                from app.ai_logger import update_ai_log, commit_ai_log
                update_ai_log(
                    response_length=150,
                    citations_count=0
                )
                commit_ai_log(success=False, error_message=f"Missing explicit references: {still_missing}")
            except:
                pass
            return

        # ── Stage 6: Context assembly (pure Python, no blocking I/O) ─────────────
        citation_block   = build_context(chunks, is_draft=_is_draft)
        compressed_block = compress_context(chunks, question, is_draft=_is_draft)

        full_rag_context = (
            (f"--- CHAT HISTORY ---\n{history_context}\n--- END HISTORY ---\n\n" if history_context else "")
            + citation_block
            + (f"\n\n{calc_block}" if calc_block else "")
            + "\n\n--- COMPRESSED STATUTORY EXCERPTS (for quick reference) ---\n\n"
            + compressed_block
        )

        from app.generation.rules_engine import rules_engine
        truth_rules_text = rules_engine.get_all_rules_as_text()

        # ── Stage 7: Streaming LLM synthesis ──────────────────────────────────────
        if _is_draft:
            _synth_msg = "Synthesizing Advisory & Drafting..."
        elif _complexity >= 0.60:
            _synth_msg = "Synthesizing Legal Position & Precedents..."
        else:
            _synth_msg = "Synthesizing Sovereign Legal Position..."

        yield f"__STATUS__:{json.dumps({'msg': _synth_msg})}__END_STATUS__"

        context_fingerprint = compute_context_fingerprint(
            question=question,
            chunks=chunks,
            assembled_context=full_rag_context,
            provider=actual_provider,
            model=actual_model
        )

        from app.generation.synthesizer import synthesize_answer_stream
        response_stream = synthesize_answer_stream(
            question, full_rag_context, session_is_draft=_is_draft,
            provider=actual_provider, model=actual_model
        )

        async for chunk in stream_and_save(
            response_stream, session_id, question,
            chunks=chunks, context=citation_block, truth_rules_text=truth_rules_text,
            query_vec=query_vec, is_draft_session=_is_draft,
            answer_provider=actual_provider, answer_model=actual_model,
            context_fingerprint=context_fingerprint
        ):
            yield chunk

    async def log_wrapper(generator):
        success = False
        error_msg = None
        status_code = 200
        try:
            async for chunk in generator:
                yield chunk
            success = True
        except Exception as e:
            success = False
            error_msg = str(e)
            status_code = 500
            raise e
        finally:
            try:
                from app.ai_logger import commit_ai_log
                commit_ai_log(success=success, http_status=status_code, error_message=error_msg)
            except Exception:
                pass

    from fastapi.responses import StreamingResponse
    return StreamingResponse(log_wrapper(rag_pipeline_orchestrator()), media_type="text/event-stream")


def log_analytics_sync(endpoint_name: str):
    def decorator(func):
        import functools
        @functools.wraps(func)
        async def wrapper(request: Request, req: QuestionRequest, *args, **kwargs):
            question = req.question.strip()
            session_id = req.session_id

            user_id, username = _get_user_info_from_req(request)
            request_id = getattr(request.state, "query_id", str(uuid.uuid4())[:8])

            from app.ai_logger import init_ai_log, commit_ai_log
            init_ai_log(
                user_id=user_id,
                username=username,
                query=question,
                endpoint=endpoint_name,
                session_id=session_id,
                request_id=request_id,
                client_ip=request.client.host if request.client else None
            )

            success = False
            error_msg = None
            status_code = 200
            try:
                res = await func(request, req, *args, **kwargs)
                success = True
                return res
            except Exception as e:
                success = False
                error_msg = str(e)
                status_code = 500
                raise e
            finally:
                try:
                    commit_ai_log(success=success, http_status=status_code, error_message=error_msg)
                except Exception:
                    pass
        return wrapper
    return decorator


@app.post("/ask-sync")
@limiter.limit("20/minute")
@log_analytics_sync("/api/ask-sync")
async def ask_question_sync(request: Request, req: QuestionRequest):
    """Non-streaming version of /ask — returns complete JSON response.
    Required for AWS API Gateway HTTP_PROXY compatibility (no SSE streaming support)."""
    import time
    t_total_start = time.monotonic()

    import asyncio as _asyncio
    import urllib.parse as _urlparse
    from fastapi.responses import JSONResponse as _JSONResponse
    from app.routing.router import route_query as _route_query
    from app.generation.synthesizer import _estimate_complexity, synthesize_answer_stream as _synth_stream, input_tokens_var, output_tokens_var
    from app.generation.calculation_engine import detect_and_calculate, format_for_context
    from app.generation.context_compressor import compress_context
    from app.cache import cache_lookup
    from app.retrieval.retriever import embed_query

    input_tokens_var.set(None)
    output_tokens_var.set(None)

    question = req.question.strip()
    session_id = req.session_id

    # Initialize timing instrumentation metrics
    cache_lookup_ms = 0.0
    query_classification_ms = 0.0
    context_build_ms = 0.0
    context_compression_ms = 0.0
    prompt_build_ms = 0.0
    generation_ms = 0.0
    validation_ms = 0.0
    repair_generation_ms = 0.0
    cache_write_ms = 0.0

    # 1. Validate requested provider & model (Fix 5)
    actual_provider, actual_model = validate_provider_and_model(req.provider, req.model)

    from app.ai_logger import update_ai_log, commit_ai_log
    update_ai_log(
        selected_provider=req.provider or "default",
        selected_model=req.model or "default",
        provider=actual_provider,
        model_used=actual_model
    )

    if session_id:
        collection = get_session_collection()
        if collection is not None:
            collection.update_one(
                {"session_id": session_id},
                {"$push": {
                    "messages": {
                        "message_id": str(uuid.uuid4()),
                        "role": "user",
                        "content": question,
                        "metadata": {},
                        "citations": [],
                        "sources": [],
                        "timestamp": datetime.now()
                    }
                }}
            )

    route = _route_query(question)
    domain_paths = route.get("domain_paths", [])
    _q = question.lower()
    _complexity = _estimate_complexity(question)

    history_context = ""
    _session_is_draft = False
    if session_id:
        collection = get_session_collection()
        if collection is not None:
            _sess = collection.find_one({"session_id": session_id})
            if _sess and "messages" in _sess:
                recent = _sess["messages"][:-1][-6:]
                for msg in recent:
                    history_context += f"{msg['role'].upper()}: {msg['content']}\n"
                _DRAFT_HIST_KW = [
                    "advisory", "our understanding", "gst implications", "provide opinion",
                    "our comments", "tax position", "gst treatment", "advise on", "legal opinion",
                    "draft", "notice", "reply", "appeal", "scn", "show cause", "drc-01",
                ]
                _session_is_draft = any(k in history_context.lower() for k in _DRAFT_HIST_KW)

    query_vec = embed_query(question)

    t_cache_start = time.monotonic()
    cached_answer = cache_lookup(question, query_vec, provider=actual_provider, model=actual_model)
    cache_lookup_ms = (time.monotonic() - t_cache_start) * 1000.0
    if cached_answer:
        cached_text, cached_sources = cached_answer
        update_ai_log(
            model_used="cache",
            cache_hit=True,
            response_length=len(cached_text),
            citations_count=len(cached_sources or [])
        )
        commit_ai_log(success=True)
        return _JSONResponse({
            "answer": cached_text,
            "sources": cached_sources or [],
            "metrics": {
                "cache_lookup_ms": round(cache_lookup_ms, 2),
                "query_classification_ms": 0.0,
                "statute_retrieval_ms": 0.0,
                "faiss_ms": 0.0,
                "bm25_ms": 0.0,
                "provision_graph_ms": 0.0,
                "rerank_ms": 0.0,
                "context_build_ms": 0.0,
                "context_compression_ms": 0.0,
                "prompt_build_ms": 0.0,
                "input_tokens": {"value": len(question) // 4, "estimated": True},
                "output_tokens": {"value": len(cached_text) // 4, "estimated": True},
                "generation_ms": 0.0,
                "validation_ms": 0.0,
                "repair_generation_ms": 0.0,
                "cache_write_ms": 0.0,
                "total_latency_ms": round((time.monotonic() - t_total_start) * 1000.0, 2)
            }
        })

    calc_result = detect_and_calculate(question)
    calc_block = format_for_context(calc_result) if calc_result else ""

    _DRAFT_KW = [
        "draft", "notice", "reply", "appeal", "submission", "advisory", "scn", "show cause",
        "drc-01", "drc 01", "asmt-10", "our understanding", "gst implications",
        "provide opinion", "our comments", "tax position", "advise on",
    ]
    _is_draft = _session_is_draft or any(k in _q for k in _DRAFT_KW)

    # Skip LLM-based query expansion in sync mode — saves ~10-15s from the extra Sonnet call.
    # Use rule-based multi-query for drafts only (no LLM needed).
    # Rule-based query topic routing
    t_class_start = time.monotonic()
    from app.retrieval.query_refiner import classify_topic_rules, extract_query_topic

    topic_info = classify_topic_rules(question)
    classification_method = "rules"
    classification_model = "rules"
    cold_start_ms = 0.0

    if topic_info.get("topic") is None:
        classification_method = "llm"
        from app.config import CLAUDE_UTILITY_MODEL, LLM_MODEL
        classification_model = CLAUDE_UTILITY_MODEL if actual_provider == "anthropic" else LLM_MODEL

        from app.generation import clients
        client_was_none = False
        if actual_provider == "anthropic":
            client_was_none = clients._claude_client is None
        elif actual_provider == "openai":
            client_was_none = clients._oai_client is None

        t_init_start = time.monotonic()
        topic_info = extract_query_topic(question, provider=actual_provider, model=None)
        if client_was_none:
            cold_start_ms = (time.monotonic() - t_init_start) * 1000.0

    query_classification_ms = (time.monotonic() - t_class_start) * 1000.0

    refined_q = question
    if _is_draft:
        advanced_queries = {
            "queries": [
                refined_q,
                refined_q + " section rule act provisions conditions eligibility liability",
                refined_q + " high court supreme court judgment held ruling decision AAR",
                refined_q + " CBIC circular notification clarification instruction",
            ],
            "hyde_document": "",
            "topic": topic_info.get("topic") or "General",
            "subtopic": topic_info.get("subtopic"),
        }
    else:
        advanced_queries = {
            "queries": [question],
            "hyde_document": "",
            "topic": topic_info.get("topic") or "General",
            "subtopic": topic_info.get("subtopic")
        }

    retriever = get_retriever()
    # skip_rerank: FlashRank takes 30-50s on 100+ candidates — bypassed to fit in 29s.
    # top_k raised now that ALB is in path and Sonnet is used for all queries.
    _top_k = 20 if _is_draft else (18 if _complexity >= 0.60 else 15)
    chunks = retriever.search(
        query=refined_q,
        top_k=_top_k,
        allowed_sources=route["use_sources"],
        advanced_queries=advanced_queries,
        domain_paths=domain_paths,
        is_draft=_is_draft,
        skip_rerank=True,
    )

    # Check and escalate retrieval coverage
    chunks, coverage_satisfied, still_missing = check_and_escalate_evidence(
        question, chunks, allowed_sources=route["use_sources"],
        domain_paths=domain_paths, is_draft=_is_draft
    )

    if not coverage_satisfied:
        missing_str = ", ".join(still_missing)
        update_ai_log(
            response_length=150,
            citations_count=0
        )
        commit_ai_log(success=False, error_message=f"Missing explicit references: {still_missing}")
        return _JSONResponse({
            "answer": f"I cannot sufficiently answer your query because the required statutory evidence for the following provision(s) is missing from the database: **{missing_str}**. Please review the source documents manually.",
            "sources": []
        })

    t_ctx_build_start = time.monotonic()
    citation_block = build_context(chunks, is_draft=_is_draft)
    context_build_ms = (time.monotonic() - t_ctx_build_start) * 1000.0

    t_ctx_comp_start = time.monotonic()
    compressed_block = compress_context(chunks, question, is_draft=_is_draft)
    context_compression_ms = (time.monotonic() - t_ctx_comp_start) * 1000.0

    t_prompt_start = time.monotonic()
    full_rag_context = (
        (f"--- CHAT HISTORY ---\n{history_context}\n--- END HISTORY ---\n\n" if history_context else "")
        + citation_block
        + (f"\n\n{calc_block}" if calc_block else "")
        + "\n\n--- COMPRESSED STATUTORY EXCERPTS (for quick reference) ---\n\n"
        + compressed_block
    )
    prompt_build_ms = (time.monotonic() - t_prompt_start) * 1000.0

    context_fingerprint = compute_context_fingerprint(
        question=question,
        chunks=chunks,
        assembled_context=full_rag_context,
        provider=actual_provider,
        model=actual_model
    )

    # Draft/advisory queries: Haiku + 4000 tokens (~23-27s) — fits API Gateway's 29s limit.
    # Non-draft Q&A: Sonnet (responses are 1000-2500 tokens, ~15-25s) — fits and gives full quality.
    t_gen_start = time.monotonic()

    is_strict = _is_draft or any(kw in question.lower() for kw in [
        "draft", "notice", "reply", "appeal", "submission", "advisory", "scn", "show cause"
    ])

    # Local helper to build concise disclaimers
    def _build_concise_disclaimer(val_res_obj):
        unverified_citations = [cit for cit, status in val_res_obj.get("citations_status", {}).items() if status == "UNVERIFIED"]
        ungrounded_nums = val_res_obj.get("ungrounded_numbers", [])

        disclaimer = (
            "I cannot sufficiently verify this answer from the available legal sources.\n\n"
            "**Verification Issues:**\n"
        )
        if unverified_citations:
            disclaimer += f"- Unverified Citations (absent from source files): {', '.join(f'`{c}`' for c in unverified_citations)}\n"
        if ungrounded_nums:
            disclaimer += f"- Ungrounded Statutory Parameters/Values: {', '.join(f'`{n}`' for n in ungrounded_nums)}\n"

        other_warnings = []
        for w in val_res_obj.get("warnings", []):
            if "Unverified Claim" in w:
                import re
                fn_match = re.search(r"in text of '([^']+)'", w)
                fn = fn_match.group(1) if fn_match else "source file"
                other_warnings.append(f"Claim under citation could not be verified in `{fn}`")
            elif "Contradiction" in w:
                import re
                fn_match = re.search(r"contradicts source '([^']+)'", w)
                fn = fn_match.group(1) if fn_match else "source file"
                other_warnings.append(f"Statement contradicts information in `{fn}`")
            elif "Authority Mismatch" in w:
                other_warnings.append(w)

        if other_warnings:
            for ow in sorted(list(set(other_warnings))):
                disclaimer += f"- {ow}\n"

        disclaimer += "\nPlease review the source documents manually or refine your query to match the available context."
        return disclaimer

    usage_initial = {}
    usage_repair = {}
    usage_standard = {}

    if is_strict:
        t_gen_only_start = time.monotonic()
        def _gen_initial():
            return "".join(_synth_stream(question, full_rag_context, session_is_draft=_is_draft, force_haiku=_is_draft, provider=actual_provider, model=actual_model, usage_tracker=usage_initial))
        initial_answer = await _asyncio.to_thread(_gen_initial)
        generation_ms = (time.monotonic() - t_gen_only_start) * 1000.0

        t_val_start = time.monotonic()
        from app.generation.validator import validate_answer_integrity
        val_res = validate_answer_integrity(initial_answer, chunks or [], is_strict=True)
        validation_ms = (time.monotonic() - t_val_start) * 1000.0

        repair_generation_ms = 0.0
        if val_res["is_valid"]:
            answer = initial_answer
        else:
            # Check if the validation failure is repairable from existing evidence.
            # It is NOT repairable if the evidence lacks the cited provisions/citations,
            # or lacks the statutory parameters/numbers cited in the warnings.
            is_repairable = True

            unverified_cits = [cit for cit, status in val_res.get("citations_status", {}).items() if status == "UNVERIFIED"]
            if unverified_cits:
                is_repairable = False

            if val_res.get("ungrounded_numbers"):
                is_repairable = False

            if is_repairable:
                warnings_str = "\n".join(f"- {w}" for w in val_res["warnings"])
                correction_prompt = f"""
                [COMPLIANCE CORRECTION REQUIRED]
                Your previous answer failed legal compliance checks with these errors:
                {warnings_str}

                Please regenerate the response correcting these errors. Only cite valid facts from sources:
                Sources:
                {full_rag_context}

                Previous Answer:
                {initial_answer}
                """

                t_repair_start = time.monotonic()
                def _gen_repair():
                    return "".join(_synth_stream(correction_prompt, full_rag_context, session_is_draft=_is_draft, force_haiku=_is_draft, provider=actual_provider, model=actual_model, usage_tracker=usage_repair))
                repaired_answer = await _asyncio.to_thread(_gen_repair)
                repair_generation_ms = (time.monotonic() - t_repair_start) * 1000.0

                t_val_repair_start = time.monotonic()
                val_res_repaired = validate_answer_integrity(repaired_answer, chunks or [], is_strict=True)
                validation_ms += (time.monotonic() - t_val_repair_start) * 1000.0
                if val_res_repaired["is_valid"]:
                    answer = repaired_answer
                else:
                    answer = _build_concise_disclaimer(val_res_repaired)
            else:
                # Fail closed immediately
                answer = _build_concise_disclaimer(val_res)
    else:
        t_gen_only_start = time.monotonic()
        def _gen_standard():
            return "".join(_synth_stream(question, full_rag_context, session_is_draft=_is_draft, force_haiku=_is_draft, provider=actual_provider, model=actual_model, usage_tracker=usage_standard))
        answer = await _asyncio.to_thread(_gen_standard)
        generation_ms = (time.monotonic() - t_gen_only_start) * 1000.0

        t_val_start = time.monotonic()
        validation_ms = 0.0
        repair_generation_ms = 0.0
        try:
            from app.generation.validator import validate_answer_integrity, validate_logic
            val_res = validate_answer_integrity(answer, chunks or [])
            validation_ms = (time.monotonic() - t_val_start) * 1000.0
            all_warnings = val_res["warnings"] + validate_logic(answer)
            if all_warnings:
                all_warnings = list(set(all_warnings))
                warning_msg = "\n\n> [!WARNING] **AUTOMATED COMPLIANCE CHECK**\n"
                warning_msg += "> The following potential issues were detected in this drafted opinion:\n"
                for w in all_warnings:
                    warning_msg += f"> - {w}\n"
                warning_msg += "> \n> *Please verify these points manually before professional use.*"
                answer += warning_msg
        except Exception as val_err:
            logger.warning(f"Post-gen validation failed: {val_err}")

    t_gen_end = time.monotonic()

    unique_sources: list = []
    seen_src: set = set()
    for c in chunks:
        _dedup_rel = c.get("rel_path", "") or c.get("metadata", {}).get("rel_path", "") or c.get("source", "")
        src_key = (_dedup_rel, c.get("page", 0))
        if src_key not in seen_src:
            seen_src.add(src_key)
            _raw_src = c.get("source", "") or c.get("metadata", {}).get("source", "")
            _rel_path = c.get("rel_path", "") or c.get("metadata", {}).get("rel_path", "")
            _basename = os.path.basename(_rel_path) if _rel_path else os.path.basename(_raw_src)
            _enc_path = _urlparse.quote(_rel_path.replace("\\", "/"), safe="") if _rel_path else ""
            _enc_name = _urlparse.quote(_basename, safe="")
            _url = (
                f"/api/documents/view_by_path?path={_enc_path}"
                if _enc_path else
                f"/api/documents/view?category=all&filename={_enc_name}"
            )
            _snippet = (c.get("text") or "").strip()
            unique_sources.append({
                "title": _basename or "Document",
                "page": c.get("page", 1),
                "url": _url,
                "rel_path": _rel_path,
                "score": float(c.get("_rerank_score", 0)),
                "snippet": _snippet[:800] if _snippet else "",
            })
        if len(unique_sources) >= 8:
            break

    if session_id and answer.strip():
        collection = get_session_collection()
        if collection is not None:
            collection.update_one(
                {"session_id": session_id},
                {
                    "$push": {
                        "messages": {
                            "message_id": str(uuid.uuid4()),
                            "role": "assistant",
                            "content": answer,
                            "metadata": {},
                            "citations": unique_sources,
                            "sources": unique_sources,
                            "timestamp": datetime.now()
                        }
                    },
                    "$set": {"updated_at": datetime.now()},
                }
            )

    # 4-layer/Confidence-gated cache store for sync requests
    t_cache_write_start = time.monotonic()
    is_valid = True
    if 'val_res_repaired' in locals():
        is_valid = val_res_repaired["is_valid"]
    elif 'val_res' in locals():
        is_valid = val_res["is_valid"]

    if answer.strip() and query_vec is not None and is_valid:
        try:
            from app.cache import cache_store
            confidence = 0.95 if is_strict else 0.80
            cache_store(question, query_vec, answer, confidence, unique_sources if chunks else [])
        except Exception as ce:
            logger.warning(f"Cache store error in ask-sync (non-fatal): {ce}")
    cache_write_ms = (time.monotonic() - t_cache_write_start) * 1000.0

    # Retrieve actual token counts mutably set inside worker thread
    if is_strict:
        provider_input_tokens = usage_initial.get("input_tokens", 0) + usage_repair.get("input_tokens", 0)
        provider_output_tokens = usage_initial.get("output_tokens", 0) + usage_repair.get("output_tokens", 0)
    else:
        provider_input_tokens = usage_standard.get("input_tokens", 0)
        provider_output_tokens = usage_standard.get("output_tokens", 0)

    # Commit AI log analytics
    try:
        from app.ai_logger import update_ai_log, commit_ai_log
        update_ai_log(
            generation_time_ms=round((t_gen_end - t_gen_start) * 1000, 2),
            estimated_completion_tokens=len(answer) // 4,
            response_length=len(answer),
            citations_count=len(unique_sources)
        )
        commit_ai_log(success=bool(answer.strip()))
    except Exception as cle:
        logger.warning(f"AI logger commit error (non-fatal): {cle}")

    input_tokens_data = {"value": provider_input_tokens, "estimated": False}
    output_tokens_data = {"value": provider_output_tokens, "estimated": False}

    metrics = {
        "query_classification_ms": round(query_classification_ms, 2),
        "statute_retrieval_ms": round(getattr(chunks, "metrics", {}).get("statute_retrieval_ms", 0.0), 2),
        "faiss_ms": round(getattr(chunks, "metrics", {}).get("faiss_ms", 0.0), 2),
        "bm25_ms": round(getattr(chunks, "metrics", {}).get("bm25_ms", 0.0), 2),
        "provision_graph_ms": round(getattr(chunks, "metrics", {}).get("provision_graph_ms", 0.0), 2),
        "rerank_ms": round(getattr(chunks, "metrics", {}).get("rerank_ms", 0.0), 2),
        "cache_lookup_ms": round(cache_lookup_ms, 2),
        "context_build_ms": round(context_build_ms, 2),
        "context_compression_ms": round(context_compression_ms, 2),
        "prompt_build_ms": round(prompt_build_ms, 2),
        "provider_input_tokens": provider_input_tokens,
        "provider_output_tokens": provider_output_tokens,
        "final_answer_character_count": len(answer),
        "estimated_visible_answer_tokens": len(answer) // 4,
        "input_tokens": input_tokens_data,
        "output_tokens": output_tokens_data,
        "generation_ms": round(generation_ms, 2),
        "validation_ms": round(validation_ms, 2),
        "repair_generation_ms": round(repair_generation_ms, 2),
        "cache_write_ms": round(cache_write_ms, 2),
        "total_latency_ms": round((time.monotonic() - t_total_start) * 1000.0, 2),
        "classification_method": classification_method,
        "classification_model": classification_model,
        "classification_cold_start_ms": round(cold_start_ms, 2),
    }

    val_report = None
    if 'val_res_repaired' in locals() and not val_res["is_valid"]:
        val_report = val_res_repaired
    elif 'val_res' in locals():
        val_report = val_res

    # Clean retrieved_chunks for JSON serialization
    serialized_chunks = []
    for c in chunks:
        c_copy = c.copy()
        if "_matched_provisions" in c_copy:
            c_copy.pop("_matched_provisions", None)
        serialized_chunks.append(c_copy)

    return _JSONResponse({
        "answer": answer,
        "sources": unique_sources,
        "retrieved_chunks": serialized_chunks,
        "validation_report": val_report,
        "metrics": metrics,
        "answer_provider": actual_provider,
        "answer_model": actual_model,
        "context_fingerprint": context_fingerprint
    })


async def _execute_ask_question_with_file(
    request: Request,
    file: UploadFile,
    question_text: str,
    session_id: Optional[str],
    provider: str = None,
    model: str = None,
):

    # 1. Read the file
    file_bytes = await file.read()
    filename = (file.filename or "").lower()

    extracted_text = ""
    # 2. Parse based on extension
    if filename.endswith(".pdf"):
        from app.ingestion.pdf_scanned import extract_text_from_scanned_pdf
        import tempfile
        import os
        # Need to save to temp file for fitz to open
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            pages = extract_text_from_scanned_pdf(tmp_path)
            extracted_text = "\n".join(p["text"] for p in pages)
        finally:
            os.remove(tmp_path)

    elif filename.endswith((".png", ".jpg", ".jpeg")):
        from app.ingestion.pdf_scanned import extract_text_from_image
        extracted_text = extract_text_from_image(file_bytes)
    elif filename.endswith(".txt"):
        extracted_text = file_bytes.decode("utf-8", errors="replace")
    elif filename.endswith(".docx"):
        try:
            import docx as _docx
            from io import BytesIO as _BytesIO
            _doc = _docx.Document(_BytesIO(file_bytes))
            _paras = [p.text for p in _doc.paragraphs if p.text.strip()]
            # Also extract table cell text
            for _tbl in _doc.tables:
                for _row in _tbl.rows:
                    for _cell in _row.cells:
                        if _cell.text.strip():
                            _paras.append(_cell.text.strip())
            extracted_text = "\n".join(_paras)
        except Exception as _e:
            extracted_text = f"[DOCX extraction error: {_e}]"
    else:
        extracted_text = "[Unsupported file format. Please upload PDF, DOCX, TXT, PNG, JPG, or JPEG.]"

    # Save User Question
    if session_id:
        collection = get_session_collection()
        if collection is not None:
             collection.update_one(
                {"session_id": session_id},
                {"$push": {
                    "messages": {
                        "message_id": str(uuid.uuid4()),
                        "role": "user",
                        "content": question_text,
                        "metadata": {},
                        "citations": [],
                        "sources": [],
                        "timestamp": datetime.now()
                    }
                }}
            )

    # Fetch history
    history_context = ""
    if session_id:
        collection = get_session_collection()
        if collection is not None:
            session = collection.find_one({"session_id": session_id})
            if session and "messages" in session:
                recent = session["messages"][:-1][-6:]
                for msg in recent:
                    history_context += f"{msg['role'].upper()}: {msg['content']}\n"

    # Route + optional query expansion (blocked Haiku call — skip for simple queries)
    route = route_query(question_text)
    from app.generation.synthesizer import _estimate_complexity
    _file_complexity = _estimate_complexity(question_text)
    if _file_complexity >= 0.80:
        from app.retrieval.query_refiner import generate_advanced_queries
        advanced_queries = generate_advanced_queries(question_text, provider=provider, model=model)
        refined_q = advanced_queries.get("queries", [question_text])[0]
    else:
        advanced_queries = {"queries": [question_text], "hyde_document": "", "topic": "General", "subtopic": None}
        refined_q = question_text

    retriever = get_retriever()
    chunks = retriever.search(
        query=refined_q,
        top_k=15,
        allowed_sources=route["use_sources"],
        advanced_queries=advanced_queries,
        domain_paths=route.get("domain_paths", []),
    )
    # Check and escalate retrieval coverage
    chunks, coverage_satisfied, still_missing = check_and_escalate_evidence(
        question_text, chunks, allowed_sources=route["use_sources"],
        domain_paths=route.get("domain_paths", []), is_draft=False
    )

    if not coverage_satisfied:
        missing_str = ", ".join(still_missing)
        async def err_generator():
            yield f"__STATUS__:{json.dumps({'msg': 'Retrieval Escalation failed to recover required evidence.'})}__END_STATUS__"
            yield f"\nI cannot sufficiently answer your query because the required statutory evidence for the following provision(s) is missing from the database: **{missing_str}**. Please review the source documents manually."
            try:
                from app.ai_logger import update_ai_log, commit_ai_log
                update_ai_log(
                    response_length=150,
                    citations_count=0
                )
                commit_ai_log(success=False, error_message=f"Missing explicit references: {still_missing}")
            except:
                pass
        return StreamingResponse(err_generator(), media_type="text/event-stream")

    from app.generation.context_builder import build_context
    from app.generation.context_compressor import compress_context
    rag_context = (
        build_context(chunks)
        + "\n\n--- COMPRESSED STATUTORY EXCERPTS (for quick reference) ---\n\n"
        + compress_context(chunks, question_text)
    )

    # Combine File Content with RAG Context
    file_context = ""
    if extracted_text.strip():
        file_context = f"\n--- UPLOADED FILE CONTENT ({file.filename}) ---\n{extracted_text}\n--- END UPLOADED FILE ---\n\n"

    full_rag_context = file_context + rag_context
    if history_context:
        full_rag_context = f"--- CHAT HISTORY ---\n{history_context}\n--- END HISTORY ---\n\n" + full_rag_context

    # Build truth rules text for hallucination guard
    from app.generation.rules_engine import rules_engine
    truth_rules_text = rules_engine.get_all_rules_as_text()

    # Generate Stream
    from app.generation.synthesizer import synthesize_answer_stream
    response_stream = synthesize_answer_stream(
        question_text, full_rag_context,
        provider=provider, model=model
    )

    from fastapi.responses import StreamingResponse
    wrapped_stream = stream_and_save(
        response_stream, session_id, question_text,
        chunks=chunks, context=rag_context, truth_rules_text=truth_rules_text,
        answer_provider=provider, answer_model=model
    )

    return StreamingResponse(wrapped_stream, media_type="text/event-stream")


@app.post("/ask-with-file")
@limiter.limit("20/minute")
async def ask_question_with_file(
    request: Request,
    file: UploadFile = File(...),
    question: str = Form(...),
    session_id: Optional[str] = Form(None),
    provider: Optional[str] = Form(None),
    model: Optional[str] = Form(None),
):
    question_text = question.strip()

    # 1. Validate requested provider & model (Fix 5)
    actual_provider, actual_model = validate_provider_and_model(provider, model)

    user_id, username = _get_user_info_from_req(request)
    request_id = getattr(request.state, "query_id", str(uuid.uuid4())[:8])

    from app.ai_logger import init_ai_log
    init_ai_log(
        user_id=user_id,
        username=username,
        query=question_text,
        endpoint="/api/ask-with-file",
        session_id=session_id,
        request_id=request_id,
        client_ip=request.client.host if request.client else None
    )
    from app.ai_logger import update_ai_log
    update_ai_log(
        selected_provider=provider or "default",
        selected_model=model or "default",
        provider=actual_provider,
        model_used=actual_model
    )

    try:
        return await _execute_ask_question_with_file(request, file, question_text, session_id, provider=actual_provider, model=actual_model)
    except Exception as e:
        try:
            from app.ai_logger import commit_ai_log
            commit_ai_log(success=False, http_status=500, error_message=str(e))
        except Exception:
            pass
        raise e


# ---------- Feedback Endpoint ----------

class FeedbackRequest(BaseModel):
    session_id: Optional[str] = None
    question: str
    answer_preview: str
    rating: int          # 1 = thumbs up, -1 = thumbs down
    comment: Optional[str] = None

@app.post("/feedback")
@limiter.limit("60/minute")
async def submit_feedback(request: Request, req: FeedbackRequest):
    """Stores user ratings (thumbs up/down) for quality monitoring and future fine-tuning."""
    import logging
    _fb_logger = logging.getLogger("feedback")
    try:
        from app.database import get_db
        db = get_db()
        if db is not None:
            db["feedback"].insert_one({
                "session_id": req.session_id,
                "question": req.question,
                "answer_preview": req.answer_preview[:200],
                "rating": req.rating,
                "comment": req.comment,
                "user": "anonymous",
                "timestamp": datetime.now(),
            })
        _fb_logger.info(f"Feedback recorded | rating={req.rating} | q={req.question[:60]}")
        return {"status": "recorded", "rating": req.rating}
    except Exception as e:
        _fb_logger.error(f"Feedback save error: {e}")
        return {"status": "error", "detail": str(e)}


# ---------- PDF Reporting Endpoint ----------
from fastapi.responses import Response
from app.generation.pdf_report import PDFReportGenerator
import os

pdf_gen = PDFReportGenerator(output_dir="data/generated_reports")

class PDFRequest(BaseModel):
    title: str
    content: str

@app.post("/generate-pdf")
def create_pdf(req: PDFRequest):
    # Use the class to generate PDF
    # We construct a filename
    import hashlib
    title_hash = hashlib.md5(req.title.encode()).hexdigest()[:8]
    filename = f"Report_{title_hash}.pdf"
    pdf_path = pdf_gen.generate_report(req.content, filename)

    if not os.path.exists(pdf_path):
        return Response(status_code=500, content="Error generating PDF")

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=advisory_report.pdf"}
    )

# ---------- Frontend Serving ----------
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

# We look for the dist directory which should exist in the multi-stage docker container
frontend_dist_path = Path(__file__).resolve().parent.parent.parent / "frontend" / "dist"

if frontend_dist_path.exists() and frontend_dist_path.is_dir():
    # Mount assets directly
    assets_path = frontend_dist_path / "assets"
    if assets_path.exists():
        app.mount("/assets", StaticFiles(directory=assets_path), name="assets")

    # Catch-all route for SPA
    @app.get("/{full_path:path}")
    async def serve_frontend(full_path: str):
        # Ignore API routes
        if full_path.startswith("api/"):
            return Response(status_code=404)

        filepath = frontend_dist_path / full_path
        if filepath.exists() and filepath.is_file():
            return FileResponse(filepath)

        # Fallback to index.html for React Router
        index_path = frontend_dist_path / "index.html"
        if index_path.exists():
            return FileResponse(index_path)

        return Response(status_code=404, content="Frontend not found")

# Trigger reload comment 2
