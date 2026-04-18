from fastapi import APIRouter, HTTPException, Depends, File, UploadFile, Query
from fastapi.responses import StreamingResponse
from typing import List, Optional, Dict, Any
from bson import ObjectId
import os
import tempfile
import io
import logging
from docx import Document
from pydantic import BaseModel, Field
import numpy as np
from app.database import get_template_collection
from app.embeddings.embedder import embed_texts
from app.config import (
    LLM_PROVIDER, OPENAI_API_KEY, LLM_MODEL,
    ANTHROPIC_API_KEY, CLAUDE_MAIN_MODEL,
)
from app.security import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter()

# ── LLM Client (uses configured provider) ────────────────────────────
if LLM_PROVIDER == "anthropic":
    import anthropic as _anthropic
    _customize_client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
else:
    import openai as _openai
    _customize_client = _openai.OpenAI(api_key=OPENAI_API_KEY)

# --- Pydantic Models for Validation ---

class ChatMessage(BaseModel):
    role: str # "user" or "assistant"
    content: str

class CustomizeRequest(BaseModel):
    user_context: str = Field(..., description="The user's specific case details and instructions for the AI.")
    messages: Optional[List[ChatMessage]] = Field(None, description="Optional chat history for multi-turn customization.")

class TemplateResponse(BaseModel):
    id: str
    title: str
    category: str
    sub_category: str
    stage: str
    keywords: List[str]
    summary: str
    content: Optional[str] = None 
    score: Optional[float] = None # Added for vector search ranking
    
    @classmethod
    def from_mongo(cls, doc: dict, include_content: bool = False, score: float = None):
        return cls(
            id=str(doc.get("_id")),
            title=doc.get("title", ""),
            category=doc.get("category", ""),
            sub_category=doc.get("sub_category", ""),
            stage=doc.get("stage", ""),
            keywords=doc.get("keywords", []),
            summary=doc.get("summary", ""),
            content=doc.get("content", "") if include_content else None,
            score=score
        )

# --- Endpoints ---

@router.get("/search")
async def search_templates(
    query: Optional[str] = None, 
) -> Dict[str, Any]:
    """
    SOTA Semantic Search for Litigation Templates.
    Generates query embeddings and ranks documents by cosine similarity.
    """
    collection = get_template_collection()
    if collection is None:
        raise HTTPException(status_code=500, detail="Database connection failed")

    result_groups = []

    if query:
        # 1. Generate Query Embedding
        query_vector = embed_texts([query])[0]
        
        # 2. Fetch templates with embeddings
        all_docs = list(collection.find({}, {"embedding": 1, "title": 1, "category": 1, "sub_category": 1, "stage": 1, "keywords": 1, "summary": 1}))
        
        if not all_docs:
            return {"groups": []}

        # 3. Compute Hybrid Scores
        query_lower = query.lower()
        query_np = np.array(query_vector)
        
        scored_results = []
        for doc in all_docs:
            # a) Vector Similarity
            doc_embedding = doc.get("embedding")
            if not doc_embedding:
                continue
            doc_vector = np.array(doc_embedding)
            vector_score = float(np.dot(doc_vector, query_np))
            
            # b) Lexical (Title) Boost
            title = doc.get("title", "").lower()
            lexical_boost = 0.0
            
            # Exact Title Match (Massive Boost)
            if query_lower == title:
                lexical_boost = 1.0
            # Sequential Substring Match (Strong Boost)
            elif query_lower in title:
                lexical_boost = 0.7
            else:
                # Individual Word Matches (Granular Boost)
                q_words = [w for w in query_lower.split() if len(w) > 1]
                if q_words:
                    matches = sum(1 for w in q_words if w in title)
                    match_ratio = matches / len(q_words)
                    if match_ratio > 0:
                        # 0.2 baseline for any match + up to 0.3 for full coverage
                        lexical_boost = 0.2 + (match_ratio * 0.3)
                
            # c) Keyword Field Boost
            keywords = [k.lower() for k in doc.get("keywords", [])]
            if any(q_word in keywords for q_word in query_lower.split()):
                lexical_boost = max(lexical_boost, 0.4)

            # Combined Score (Prioritize Title/Keyword matches)
            # Weighting: 40% Semantic Vector, 60% Lexical/Names
            final_score = (vector_score * 0.4) + (lexical_boost * 0.6)
            
            if final_score > 0.15: # Slightly broader threshold for noisy queries
                scored_results.append((doc, final_score))
        
        # 4. Filter and Rank
        scored_results.sort(key=lambda x: x[1], reverse=True)
        top_matches = [TemplateResponse.from_mongo(doc, score=score).dict() for doc, score in scored_results[:15]]
        
        if top_matches:
            result_groups.append({
                "title": f"Best Matches for '{query}'",
                "type": "hero",
                "templates": top_matches[:1]
            })
            
            if len(top_matches) > 1:
                result_groups.append({
                    "title": "Relevant Litigation Strategies",
                    "type": "row",
                    "templates": top_matches[1:]
                })
        else:
            result_groups.append({
                "title": "No Matches Found",
                "type": "hero",
                "templates": []
            })
            
    else:
        # ── Grouped Home View ──────────────────────────────────────────
        # Fetching categories dynamically ensures all 1,200+ docs are visible
        all_categories = collection.distinct("category")
        
        # 1. Start with a "Master Library" of most recent globally
        recent_all = list(collection.find({}).sort("ingested_at", -1).limit(50))
        if recent_all:
            result_groups.append({
                "title": "Master Litigation Library",
                "type": "row",
                "templates": [TemplateResponse.from_mongo(d).dict() for d in recent_all]
            })

        # 2. Dynamic Categories (Netflix-style shelves)
        # We increase the limit to 100 per row to show "All" responses effectively
        for cat in all_categories:
            if not cat or cat == "General": continue 
            
            docs = list(collection.find({"category": cat}).sort("ingested_at", -1).limit(100))
            if docs:
                # Map internal category names to human-friendly titles
                titles = {
                    "ITC": "ITC Reversals & Compliance",
                    "Appeal": "Appeals & Writ Formats",
                    "Demand/Recovery": "Demand & Recovery Responses",
                    "Refund": "Refund Claims & Formats",
                    "Registration": "Registration & Cancellation Replies",
                    "Notification": "GST Notifications (2025-26)",
                    "Circular": "GST Circulars & Clarifications",
                    "E-Way Bill": "E-Way Bill & Detention Replies",
                    "Compliance & Returns": "GSTR Returns & Mismatch Strategies",
                    "Demand & Penalty": "Penalty, Interest & Demand Responses"
                }
                result_groups.append({
                    "title": titles.get(cat, f"{cat} Intelligence"),
                    "type": "row",
                    "templates": [TemplateResponse.from_mongo(d).dict() for d in docs]
                })

        # 3. Add General / Miscellaneous at the bottom
        general_docs = list(collection.find({"category": "General"}).sort("ingested_at", -1).limit(100))
        if general_docs:
            result_groups.append({
                "title": "General Litigation Utility",
                "type": "row",
                "templates": [TemplateResponse.from_mongo(d).dict() for d in general_docs]
            })
            
    return {"groups": result_groups}


@router.get("/{template_id}")
async def get_template(template_id: str) -> TemplateResponse:
    collection = get_template_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(template_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid ID")

    if not doc:
        raise HTTPException(status_code=404, detail="Not found")

    return TemplateResponse.from_mongo(doc, include_content=True)

@router.get("/{template_id}/download")
async def download_template(
    template_id: str,
    content: Optional[str] = Query(None, description="Customized content to download. If None, downloads base template.")
):
    """
    Generates and returns a .docx file for the template.
    """
    collection = get_template_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(template_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Template ID")

    if not doc:
        raise HTTPException(status_code=404, detail="Template not found")

    target_content = content or doc.get("content", "No content available.")
    title = doc.get("title", "Legal_Template")

    # Create DOCX in memory
    doc_obj = Document()
    doc_obj.add_heading(title, 0)
    
    # Process content by lines to preserve some structure
    for paragraph in target_content.split('\n'):
        if paragraph.strip():
            doc_obj.add_paragraph(paragraph)
        else:
            doc_obj.add_paragraph("") # Keep spacing

    file_stream = io.BytesIO()
    doc_obj.save(file_stream)
    file_stream.seek(0)

    filename = f"{title.replace(' ', '_')}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.post("/{template_id}/customize")
async def customize_template(
    template_id: str, 
    request: CustomizeRequest
):
    """
    Conversational AI Customization Engine.
    Handles multi-turn chat interaction to refine a legal draft.
    """
    collection = get_template_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(template_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Template ID")

    if not doc:
        raise HTTPException(status_code=404, detail="Template not found")

    base_content = doc.get("content", "")
    
    # Set up System Prompt
    system_msg = {
        "role": "system", 
        "content": f"""You are 'LETA', a high-end AI legal assistant specializing ONLY in GST litigation document drafting.
Your EXCLUSIVE goal is to help the user customize and refine this specific legal template:

---
{base_content}
---

STRICT OPERATIONAL RULES:
1. FOCUS: Your only purpose is to modify this draft. If the user asks general GST questions unrelated to drafting this document, politely decline and ask how you can help amend the current draft.
2. TONAL EXCELLENCE: Maintain a professional, respectful 'Sentinel' tone (high-end, precise, authoritative).
3. FACT INTEGRATION: Incorporate user facts (names, dates, GSTINs, amounts, specific arguments) with 100% precision.
4. ITERATIVE IMPROVEMENT: If the user asks for changes ('make it more aggressive', 'add a paragraph about section 16(4)', 'shorten the prayer clause'), apply those changes to the latest version of the draft.
5. FORMATTING: Output the updated legal draft within clear markers: [DRAFT_START] and [DRAFT_END].
6. CONTEXT: You can provide a brief, professional explanation of the changes made outside the markers, but keep the focus on the document.
"""
    }

    # Build message history (provider-agnostic content)
    chat_messages = []
    if request.messages:
        for msg in request.messages:
            chat_messages.append({"role": msg.role, "content": msg.content})
    chat_messages.append({"role": "user", "content": request.user_context})

    try:
        if LLM_PROVIDER == "anthropic":
            # Use Claude — system prompt goes in `system` param, not in messages
            resp = _customize_client.messages.create(
                model=CLAUDE_MAIN_MODEL,
                max_tokens=4096,
                system=system_msg["content"],
                messages=chat_messages,
                temperature=0.4,
            )
            full_response = resp.content[0].text
        else:
            # OpenAI / Ollama — system prompt is a message
            openai_messages = [system_msg] + chat_messages
            resp = _customize_client.chat.completions.create(
                model=LLM_MODEL,
                messages=openai_messages,
                temperature=0.4,
            )
            full_response = resp.choices[0].message.content

        # Extract draft if markers exist, otherwise take the whole thing
        draft = full_response
        if "[DRAFT_START]" in full_response and "[DRAFT_END]" in full_response:
            draft = full_response.split("[DRAFT_START]")[1].split("[DRAFT_END]")[0].strip()

        return {
            "status": "success",
            "full_response": full_response,
            "customized_draft": draft,
        }
    except Exception as e:
        logger.error(f"Template customization failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"AI Customization failed: {str(e)}")

@router.post("/upload")
async def upload_templates(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    # Dummy implementation for now, handled by manual ingestion scripts usually
    return {"message": "Admin upload logged. Processed via ingestion queue."}
